#!/usr/bin/env python3
"""Generate the three benchmark notebooks with the scientific logic inlined in cells.

Only IO/plotting/provenance helpers are imported (scripts/bench_utils.py); all ranking, CSLS,
CRISPRi-validation, evidence-integration and verdict logic is written directly into notebook
cells so the notebooks are readable and are not thin wrappers.

Usage:  python scripts/_build_notebooks.py [01|02|03|all]
"""
import sys
from pathlib import Path

import nbformat as nbf

ROOT = Path(__file__).resolve().parents[1]
NB = ROOT / "notebooks"
NB.mkdir(exist_ok=True)


def _nb(cells):
    nb = nbf.v4.new_notebook()
    nb.cells = cells
    nb.metadata = {
        "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
        "language_info": {"name": "python"},
    }
    return nb


def md(s):
    return nbf.v4.new_markdown_cell(s.strip("\n"))


def code(s):
    return nbf.v4.new_code_cell(s.strip("\n"))


# ============================================================ NOTEBOOK 03
def build_nb3():
    c = []
    c.append(md("""
# 03 — Benchmark: BTLA-response TF prioritisation (GREmLN masked-prior vs GENIE3 masked-raw)

Definitive head-to-head. Both models are restricted to the **common GREmLN∩GENIE3 TF universe**
and scored against the **BTLA_TCR_vs_TCR 4 h panel (249 DEGs)**:

* **GREmLN** — CSLS (k=10): a TF's score = number of BTLA seeds among its top-100 embedding
  neighbours (dense rank).
* **GENIE3** — summed **outgoing** TF→BTLA-DEG edge weight (dense rank).

**Primary nomination comparison = seed-EXCLUDED** (BTLA panel genes removed → candidate
regulators). A **seed-INCLUSIVE** ranking is reported separately as biological context, with panel
genes flagged (★) and never called discoveries.

Candidates are then examined with three **orthogonal, non-predictive** validators:
1. **CD4 CRISPRi Perturb-seq** (primary; Stim8hr) — functional target-response;
2. **Paperclip** literature (identical template);
3. **BTLA multi-omics**, split into **independent orthogonal** (protein/PTM/co-IP) vs
   **derived/contextual** (transcript DE, TF/kinase activity, BIONIC).

Held-out seed-TF recovery is **supplementary only** (pending a leakage/asymmetry audit) and does
**not** enter the verdict. Requires the outputs of notebooks 01 (GENIE3 graph) and 02 (GREmLN
embeddings); set `BTLA_BENCH_DATA_ROOT` to where those + the BTLA panel/Perturb-seq live.
"""))

    c.append(code("""
import sys, json
from pathlib import Path
import numpy as np
import pandas as pd
from scipy.stats import hypergeom, rankdata, spearmanr
from statsmodels.stats.multitest import multipletests

_here = Path.cwd()
_root = next((p for p in [_here, *_here.parents] if (p / "scripts" / "bench_utils.py").exists()), _here.parent)
sys.path.insert(0, str(_root / "scripts"))
import bench_utils as bu
import yaml

CFG = yaml.safe_load((bu.repo_root() / "config" / "benchmark_config.yaml").read_text())
# ---- locked parameters (must match notebooks 01/02) ----
CSLS_K, TOPN = 10, 100
DE_PADJ = 0.05                 # CRISPRi knockdown target-response threshold (adj p only)
STIM_PRIMARY, STIM_SENS = "Stim8hr", "Stim48hr"
RANDOM_SEED = 42
CRISPRI_MARGIN = CFG["decision_rule"]["crispri_margin"]   # decision-rule margin (from config)
N_MASKED_EXPR_GENES = int(CFG["genie3"]["n_hvg"])         # masked expression universe (HVGs)

OUT = bu.repo_root() / "results"
TAB, FIG, REG = OUT / "tables", OUT / "figures", OUT / "model_registry"
for d in (TAB, FIG, REG):
    d.mkdir(parents=True, exist_ok=True)
print("DATA_ROOT:", bu.data_root())
"""))

    c.append(md("## 1. BTLA panel (assert 249), TF list, embeddings, GENIE3 graph"))
    c.append(code("""
seeds, seed_dir, seeds_df = bu.load_seeds()
assert len(seeds) == 249, f"BTLA_TCR_vs_TCR panel must be exactly 249 genes, got {len(seeds)}"
tfs = bu.load_tfs()
genes, X = bu.load_gremln_embeddings()
edges = bu.load_edges()
n_up = (seeds_df['direction'] == 'up').sum(); n_down = (seeds_df['direction'] == 'down').sum()
print(f"BTLA panel: {len(seeds)} genes (up={n_up}, down={n_down}) | TF list: {len(tfs)}")
print(f"GREmLN embedded genes: {len(genes)} | GENIE3 masked edges: {len(edges)}")
"""))

    c.append(md("""## 2. CSLS similarity and the common (scorable-by-both) TF universe

Cosine → CSLS (Conneau et al. 2018) corrects hubness: `CSLS(i,j) = 2·cos(i,j) − r_k(i) − r_k(j)`
with `r_k` the mean cosine to the k nearest neighbours."""))
    c.append(code("""
def cosine_similarity_matrix(X):
    n = np.linalg.norm(X, axis=1, keepdims=True); n = np.where(n == 0, 1.0, n)
    Xn = X / n; S = Xn @ Xn.T; np.fill_diagonal(S, -np.inf); return S

def compute_csls_matrix(S, k=CSLS_K):
    Sf = np.where(np.isfinite(S), S, -np.inf)
    rk = np.sort(Sf, axis=1)[:, -k:].mean(axis=1)
    return 2 * S - rk[:, None] - rk[None, :]

S_csls = compute_csls_matrix(cosine_similarity_matrix(X), k=CSLS_K)

gremln_tfs = set(genes) & tfs
genie3_tfs = set(edges['regulator']) & tfs
common = gremln_tfs & genie3_tfs
print(f"GREmLN TFs: {len(gremln_tfs)} | GENIE3 regulator-TFs: {len(genie3_tfs)} | common: {len(common)}")
"""))

    c.append(md("""## 3. Ranking formulae (inlined) and rankings

`score_genes_by_seed_neighbours` (GREmLN) and the outgoing-edge-weight sum (GENIE3) are computed
for both the **seed-excluded** (primary) and **seed-inclusive** (context) views, restricted to the
common universe with dense tie-aware ranks."""))
    c.append(code("""
def score_genes_by_seed_neighbours(genes, S, seed_genes, topn=TOPN, exclude_seeds=False):
    seed_in = {g for g in seed_genes if g in set(genes)}
    G = len(genes); is_seed = np.array([g in seed_in for g in genes]); K = int(is_seed.sum())
    base = K / G if G else 0.0
    dcount = np.empty(G, int); pval = np.ones(G)
    for i in range(G):
        nn = np.argpartition(-S[i], min(topn, G - 1))[:topn]
        x = int(is_seed[nn].sum()); dcount[i] = x
        nsucc = K - int(is_seed[i])
        pval[i] = hypergeom.sf(x - 1, G - 1, nsucc, topn) if nsucc >= 0 else 1.0
    out = pd.DataFrame({"gene": genes, "gremln_csls_score": dcount,
                        "fold_enrichment": (dcount / topn) / base if base > 0 else 0.0,
                        "hypergeom_p": pval, "is_seed": is_seed})
    if exclude_seeds:
        out = out.loc[~out["is_seed"]].copy()
    return out.sort_values("gremln_csls_score", ascending=False).reset_index(drop=True)

def gremln_ranking(exclude_seeds):
    r = score_genes_by_seed_neighbours(genes, S_csls, seeds, exclude_seeds=exclude_seeds)
    t = r[r["gene"].isin(common)].copy()
    t["gremln_csls_rank"] = rankdata(-t["gremln_csls_score"], method="dense").astype(int)
    t["is_BTLA_vs_TCR_seed"] = t["gene"].isin(seeds)
    t["BTLA_vs_TCR_direction"] = t["gene"].map(lambda g: f"BTLA_{seed_dir[g]}" if g in seed_dir else "not_seed")
    t["eligible_for_fair_comparison"] = (~t["is_BTLA_vs_TCR_seed"]) if exclude_seeds else True
    return t.sort_values("gremln_csls_rank").reset_index(drop=True)

def genie3_ranking(exclude_seeds):
    st = edges[edges["target"].isin(seeds)]
    inc = (st.groupby("regulator")["weight"].agg(n_seed_targets="count", genie3_score="sum")
             .reset_index().rename(columns={"regulator": "gene"}))
    # Rank over the FULL common universe: common TFs with no outgoing edge into the BTLA seed panel
    # score 0 and tie at the bottom, mirroring GREmLN which scores every common TF (incl. zero CSLS).
    # This keeps both models on an identical candidate pool for overlap/rank-correlation.
    universe = sorted(common - set(seeds)) if exclude_seeds else sorted(common)
    inc = pd.DataFrame({"gene": universe}).merge(inc, on="gene", how="left")
    inc["n_seed_targets"] = inc["n_seed_targets"].fillna(0).astype(int)
    inc["genie3_score"] = inc["genie3_score"].fillna(0.0)
    inc = inc.sort_values("genie3_score", ascending=False).reset_index(drop=True)
    inc["genie3_rank"] = rankdata(-inc["genie3_score"], method="dense").astype(int)
    inc["is_BTLA_vs_TCR_seed"] = inc["gene"].isin(seeds)
    inc["BTLA_vs_TCR_direction"] = inc["gene"].map(lambda g: f"BTLA_{seed_dir[g]}" if g in seed_dir else "not_seed")
    return inc

rk = {}
for mode, excl in (("seed_excluded", True), ("seed_inclusive", False)):
    gr, g3 = gremln_ranking(excl), genie3_ranking(excl)
    rk[mode] = {"gr": gr, "g3": g3, "gr25": list(gr.head(25)["gene"]), "g325": list(g3.head(25)["gene"])}
    gr.to_csv(TAB / f"gremln_btla_vs_tcr_{mode}_tf_ranking.csv", index=False)
    g3.to_csv(TAB / f"genie3_btla_vs_tcr_{mode}_tf_ranking.csv", index=False)
print("seed-excluded GREmLN top10:", rk["seed_excluded"]["gr25"][:10])
print("seed-excluded GENIE3 top10:", rk["seed_excluded"]["g325"][:10])
"""))

    c.append(md("## 4. Primary comparison (seed-excluded): overlap + rank correlation. Seed-inclusive as context."))
    c.append(code("""
def compare(mode):
    gr, g3 = rk[mode]["gr"], rk[mode]["g3"]
    m = gr[["gene", "gremln_csls_score", "gremln_csls_rank", "is_BTLA_vs_TCR_seed"]].merge(
        g3[["gene", "genie3_score", "genie3_rank"]], on="gene", how="outer")
    m["rank_difference_gremln_minus_genie3"] = m["gremln_csls_rank"] - m["genie3_rank"]
    both = m.dropna(subset=["gremln_csls_rank", "genie3_rank"])
    rho = spearmanr(both["gremln_csls_rank"], both["genie3_rank"])[0] if len(both) >= 3 else np.nan
    gr25, g325 = set(rk[mode]["gr25"]), set(rk[mode]["g325"])
    ov = {"mode": mode, "spearman_rank_rho": float(rho), "n_common_ranked_by_both": int(len(both)),
          "shared_top25": sorted(gr25 & g325), "gremln_only_top25": sorted(gr25 - g325),
          "genie3_only_top25": sorted(g325 - gr25)}
    m.sort_values("gremln_csls_rank").to_csv(TAB / f"gremln_genie3_common_universe_rank_comparison_{mode}.csv", index=False)
    return ov

overlap = {m: compare(m) for m in ("seed_excluded", "seed_inclusive")}
prim = overlap["seed_excluded"]
print(f"PRIMARY (seed-excluded): Spearman rho={prim['spearman_rank_rho']:.3f}; "
      f"shared top25={len(prim['shared_top25'])}: {prim['shared_top25']}")
print("GREmLN-only:", prim["gremln_only_top25"])
print("GENIE3-only:", prim["genie3_only_top25"])
"""))

    c.append(md("""## 5. CD4 CRISPRi perturbational validation (primary; seed-excluded top-25)

For each model's top-25, its predicted BTLA-DEG targets (GENIE3: graph targets ∩ common seeds;
GREmLN: top-100 CSLS neighbours ∩ common seeds; identical common seed set for both) are tested
for differential expression after that TF's knockdown
(adj p < DE_PADJ). Signed agreement flags whether confirmed targets move with/against the BTLA
programme. Both scores are **unsigned**, so target movement does not establish direction."""))
    c.append(code("""
# Both models are restricted to the IDENTICAL common BTLA seed set (the seeds scorable by BOTH
# models: present in the GREmLN gene universe AND reachable as GENIE3 graph targets). This keeps the
# predicted-target denominators symmetric so neither model benefits from broader seed coverage.
graph_nodes = set(edges["regulator"]) | set(edges["target"])
common_seed_set = (set(seeds) & set(genes)) & graph_nodes
print("common BTLA seed set (scorable by both models):", len(common_seed_set), "of", len(set(seeds)))

def genie3_targets_by_tf(tf_list):
    sub = edges[edges["regulator"].isin(tf_list) & edges["target"].isin(common_seed_set)]
    return {tf: set(g["target"].unique()) for tf, g in sub.groupby("regulator", sort=False)}

def csls_seed_targets_by_tf(tf_list):
    gi = {g: i for i, g in enumerate(genes)}; seed_in = common_seed_set; G = len(genes); out = {}
    for tf in tf_list:
        if tf not in gi:
            out[tf] = set(); continue
        nn = np.argpartition(-S_csls[gi[tf]], min(TOPN, G - 1))[:TOPN]
        out[tf] = {genes[j] for j in nn} & seed_in
    return out

def crispri_validation(union_by_model, de_stats_path):
    import anndata as ad, h5py
    all_tfs = sorted({t for lst in union_by_model.values() for t in lst})
    de = ad.read_h5ad(de_stats_path, backed="r"); de_obs = de.obs.reset_index(drop=True)
    de_gene = de.var["gene_name"].astype(str).values; n_meas = len(de_gene)
    def arm(cond):
        mask = ((de_obs["culture_condition"].astype(str) == cond)
                & (de_obs["target_contrast_gene_name"].astype(str).isin(all_tfs)))
        ri = np.where(mask.values)[0]; order = np.argsort(ri); sidx = ri[order]; inv = np.argsort(order)
        with h5py.File(de_stats_path, "r") as f:
            adjp = f["layers/adj_p_value"][sidx, :][inv]; logfc = f["layers/log_fc"][sidx, :][inv]
        tfn = de_obs.loc[ri, "target_contrast_gene_name"].astype(str).values
        ot = de_obs.loc[ri, "ontarget_significant"].values
        nc = de_obs.loc[ri, "n_cells_target"].values if "n_cells_target" in de_obs else np.full(len(ri), np.nan)
        # Deterministic one-row-per-TF map: if a TF has multiple rows for this condition, prefer an
        # on-target-significant row (QC-passing), else keep the earliest index. Avoids silent last-wins.
        pmap = {}
        for i, t in enumerate(tfn):
            if t not in pmap or (not bool(ot[pmap[t]]) and bool(ot[i])):
                pmap[t] = i
        return adjp, logfc, ot, nc, pmap
    a8, l8, o8, c8, p8 = arm(STIM_PRIMARY); a48, l48, o48, _, p48 = arm(STIM_SENS)
    def de_info(pmap, adjp, tf):
        if tf not in pmap: return None
        padj = adjp[pmap[tf], :]; hit = np.isfinite(padj) & (padj < DE_PADJ)
        return set(de_gene[hit]), int(hit.sum())
    sign = {g: (1 if seed_dir.get(g) == "up" else -1) for g in seeds}
    rows = []
    for model, top in union_by_model.items():
        pred_by = csls_seed_targets_by_tf(top) if model == "GREmLN" else genie3_targets_by_tf(top)
        for tf in top:
            pred = set(pred_by.get(tf, set())); present = tf in p8
            d8, d48 = de_info(p8, a8, tf), de_info(p48, a48, tf)
            rec = {"TF": tf, "model": model, "is_BTLA_vs_TCR_seed": tf in seeds,
                   "crispri_arm_present_8hr": bool(present), "crispri_arm_present_48hr": bool(tf in p48),
                   "ontarget_kd_significant": bool(o8[p8[tf]]) if present else np.nan,
                   "n_predicted_btla_targets": len(pred)}
            for lab, di in (("8hr", d8), ("48hr", d48)):
                if di is None or not pred:
                    rec[f"n_de_genes_crispri_{lab}"] = di[1] if di else np.nan
                    rec[f"n_confirmed_{lab}"] = np.nan; rec[f"frac_confirmed_{lab}"] = np.nan
                    rec[f"hypergeom_p_{lab}"] = np.nan; rec[f"confirmed_targets_{lab}"] = ""
                    continue
                dg, nde = di; conf = pred & dg; k = len(conf)
                rec[f"n_de_genes_crispri_{lab}"] = nde; rec[f"n_confirmed_{lab}"] = k
                rec[f"frac_confirmed_{lab}"] = k / len(pred)
                rec[f"hypergeom_p_{lab}"] = float(hypergeom.sf(k - 1, n_meas, nde, len(pred))) if k > 0 else np.nan
                rec[f"confirmed_targets_{lab}"] = "|".join(sorted(conf))
            na = no = 0
            if present and pred and d8:
                lfc = dict(zip(de_gene, l8[p8[tf], :]))
                for g in (pred & d8[0]):
                    lf = lfc.get(g, np.nan)
                    if not np.isfinite(lf) or lf == 0 or g not in sign: continue
                    na += int(np.sign(lf) == np.sign(sign[g])); no += int(np.sign(lf) != np.sign(sign[g]))
            rec["n_confirmed_kd_agrees_btla"] = na; rec["n_confirmed_kd_opposes_btla"] = no
            rec["net_signed_agreement"] = (na - no) / (na + no) if (na + no) else np.nan
            rows.append(rec)
    df = pd.DataFrame(rows); df["adj_hypergeom_p_8hr"] = np.nan
    for model in df["model"].unique():
        m = df["model"] == model; p = df.loc[m, "hypergeom_p_8hr"].values; ok = np.isfinite(p)
        if ok.sum():
            adj = np.full(len(p), np.nan); adj[ok] = multipletests(p[ok], method="fdr_bh")[1]
            df.loc[m, "adj_hypergeom_p_8hr"] = adj
    def status(r):
        if not r["crispri_arm_present_8hr"]: return "not_in_screen"
        if r["ontarget_kd_significant"] == False: return "failed_ontarget_qc"
        nconf, p, net = r["n_confirmed_8hr"], r["hypergeom_p_8hr"], r["net_signed_agreement"]
        if pd.notna(nconf) and nconf >= 3 and pd.notna(net) and net <= -0.5: return "usable_contradictory"
        if pd.notna(nconf) and nconf > 0 and pd.notna(p) and p < 0.05: return "usable_supportive"
        return "usable_unsupported"
    df["validation_status"] = df.apply(status, axis=1)
    return df

union_se = {"GREmLN": rk["seed_excluded"]["gr25"], "GENIE3": rk["seed_excluded"]["g325"]}
crispri = crispri_validation(union_se, bu.artifact("de_stats"))
crispri.to_csv(TAB / "gremln_genie3_cd4_crispri_validation.csv", index=False)

_QC_FRAC = {}   # QC-passed native fractions per model, for the difference bootstrap CI in the verdict
def crispri_stats(model):
    sub = crispri[(crispri.model == model) & (crispri.TF.isin(union_se[model]))]
    pres = sub[sub.crispri_arm_present_8hr == True]
    # PRIMARY population = perturbations passing the dataset on-target-significance flag (QC-passed).
    # This matches the canonical crispri_model_summary / Table 2 / Figure 3.
    qc = pres[pres.ontarget_kd_significant == True]
    fr = qc.frac_confirmed_8hr.astype(float).dropna()
    # SENSITIVITY population = all in-screen candidates (incl. QC failures), reported separately.
    fr_all = pres.frac_confirmed_8hr.astype(float).dropna()
    rng = np.random.default_rng(RANDOM_SEED)
    b = [rng.choice(fr.values, len(fr), replace=True).mean() for _ in range(5000)] if len(fr) else [np.nan]
    _QC_FRAC[model] = fr.values
    return {"model": model, "present": int(len(pres)),
            "qc_pass": int((sub.ontarget_kd_significant == True).sum()), "n_in_primary": int(len(fr)),
            "mean_frac_8hr": round(float(fr.mean()), 4) if len(fr) else np.nan,
            "median_frac_8hr": round(float(fr.median()), 4) if len(fr) else np.nan,
            "ci_lo": round(float(np.nanquantile(b, .025)), 4), "ci_hi": round(float(np.nanquantile(b, .975)), 4),
            "mean_frac_8hr_all_screen": round(float(fr_all.mean()), 4) if len(fr_all) else np.nan,
            **{s: int((sub.validation_status == s).sum()) for s in
               ["usable_supportive", "usable_unsupported", "usable_contradictory", "failed_ontarget_qc", "not_in_screen"]}}
cr_stats = {m: crispri_stats(m) for m in ("GREmLN", "GENIE3")}
pd.DataFrame(cr_stats.values()).to_csv(TAB / "crispri_summary_by_model_seed_excluded.csv", index=False)
pd.DataFrame(cr_stats.values())
"""))

    c.append(md("""## 6. Paperclip literature (identical template; annotation only)

Structured literature evidence for the **full union of both top-25 lists** under one identical query
template (`scripts/build_paperclip_review.py`; committed to `results/paperclip/`). Because coverage
is complete, a quantitative literature comparison across the model-specific candidate sets is drawn
here. Literature is annotation, never a predictive input."""))
    c.append(code("""
union_tfs = sorted(set(union_se["GREmLN"]) | set(union_se["GENIE3"]))
pc_path = bu.repo_root() / "results" / "paperclip" / "paperclip_union_top25_review.csv"
if pc_path.exists():
    pc = pd.read_csv(pc_path)
    pc_join = pd.DataFrame({"TF": union_tfs}).merge(pc, on="TF", how="left")
    pc_join["paperclip_reviewed"] = pc_join["paperclip_reviewed"].fillna(False).astype(bool)
else:
    pc_join = pd.DataFrame({"TF": union_tfs}); pc_join["paperclip_reviewed"] = False
    pc_join["paperclip_evidence_tier"] = np.nan
pc_join.to_csv(TAB / "paperclip_union_top25.csv", index=False)
paperclip_coverage_complete = bool(len(pc_join) and pc_join["paperclip_reviewed"].all())
strong_mod = set(pc_join.loc[pc_join["paperclip_evidence_tier"].isin(["strong", "moderate"]), "TF"])
gr_only = set(union_se["GREmLN"]) - set(union_se["GENIE3"])
g3_only = set(union_se["GENIE3"]) - set(union_se["GREmLN"])
shared = set(union_se["GREmLN"]) & set(union_se["GENIE3"])
paperclip_lit = {
    "coverage_complete": paperclip_coverage_complete,
    "shared_with_strong_moderate_lit": f"{len(shared & strong_mod)}/{len(shared)}",
    "gremln_only_with_strong_moderate_lit": f"{len(gr_only & strong_mod)}/{len(gr_only)}",
    "genie3_only_with_strong_moderate_lit": f"{len(g3_only & strong_mod)}/{len(g3_only)}",
}
print(f"Paperclip coverage: {pc_join['paperclip_reviewed'].sum()}/{len(pc_join)} union TFs; "
      f"complete={paperclip_coverage_complete}")
for k, v in paperclip_lit.items():
    print(f"  {k}: {v}")
"""))

    c.append(md("""## 7. BTLA multi-omics — independent orthogonal vs derived/contextual

Independent orthogonal = protein, phosphosite, co-IP (not derivable from the expression signal
driving the models). Derived/contextual = transcript DE, TF activity, kinase activity, BIONIC/GNN
modules. Kept **separate**; not collapsed into one count."""))
    c.append(code("""
INDEP = ["protein_support", "phosphosite_support", "coip_support"]
DERIV = ["transcript_support", "tf_activity_support", "kinase_activity_support",
         "bionic_support", "early_synapse_trafficking_support"]
mo_path = bu.artifact("multiomics_summary")
union_df = pd.DataFrame({"TF": union_tfs})
union_df["is_BTLA_vs_TCR_seed"] = union_df["TF"].isin(seeds)
union_df["BTLA_vs_TCR_direction"] = union_df["TF"].map(lambda g: f"BTLA_{seed_dir[g]}" if g in seed_dir else "not_seed")
union_df["in_gremln_top25"] = union_df["TF"].isin(union_se["GREmLN"])
union_df["in_genie3_top25"] = union_df["TF"].isin(union_se["GENIE3"])

integ = union_df.copy()
if mo_path.exists():
    mo = pd.read_csv(mo_path).rename(columns={"gene": "TF"})
    cols = ["TF"] + [c for c in INDEP + DERIV + ["strongest_support_layer", "independent_support",
            "overall_multiomics_support"] if c in mo.columns]
    integ = integ.merge(mo[cols], on="TF", how="left")
    def any_support(row, layers):
        return any(str(row.get(l, "")).lower() not in ("", "nan", "none", "no", "false", "0", "0.0")
                   for l in layers if l in integ.columns)
    integ["independent_orthogonal_support"] = integ.apply(lambda r: any_support(r, INDEP), axis=1)
    integ["derived_contextual_support"] = integ.apply(lambda r: any_support(r, DERIV), axis=1)
# attach compact CRISPRi + paperclip
cr_agg = crispri.groupby("TF").agg(crispri_status=("validation_status", lambda s: ";".join(sorted(set(s)))),
        crispri_best_frac_8hr=("frac_confirmed_8hr", "max")).reset_index()
pc_cols = [c for c in ["TF", "paperclip_evidence_tier", "paperclip_primary_phenotype",
           "paperclip_direction", "paperclip_key_source", "paperclip_short_rationale",
           "paperclip_reviewed"] if c in pc_join.columns]
integ = integ.merge(cr_agg, on="TF", how="left").merge(pc_join[pc_cols], on="TF", how="left")
integ.to_csv(TAB / "gremln_genie3_top25_integrated_evidence.csv", index=False)
print("Independent orthogonal support among union TFs:",
      int(integ.get("independent_orthogonal_support", pd.Series(dtype=bool)).sum()))
integ.head(12)
"""))

    c.append(md("## 8. Report figures (primary seed-excluded + seed-inclusive context + coverage)"))
    c.append(code("""
bu.fig_top25_comparison(rk["seed_excluded"]["gr"], rk["seed_excluded"]["g3"],
                        FIG / "fig1_top25_seed_excluded.png", seeds=seeds, mode="seed-excluded")
bu.fig_overlap(rk["seed_excluded"]["gr25"], rk["seed_excluded"]["g325"],
               FIG / "fig2_top25_overlap_seed_excluded.png", seeds=seeds, title="Seed-excluded top-25 overlap")
bu.fig_top25_comparison(rk["seed_inclusive"]["gr"], rk["seed_inclusive"]["g3"],
                        FIG / "figS_top25_seed_inclusive.png", seeds=seeds, mode="seed-inclusive")
bu.fig_crispri(crispri, FIG / "fig3_crispri_target_response.png")
bu.fig_evidence_heatmap(integ, FIG / "fig4_multiomics_heatmap.png")
n_genie3_nodes = len(set(edges['regulator']) | set(edges['target']))
bu.fig_coverage(N_MASKED_EXPR_GENES, n_genie3_nodes, len(genes),
                gremln_tfs, genie3_tfs, len(set(seeds) & common), FIG / "fig5_coverage.png")
print("figures written to", FIG)
"""))

    c.append(md("""## 9. SCENIC scope note

> SCENIC pruning of the persisted top-50k GENIE3 graph was retained as an exploratory motif-support
> sensitivity analysis (notebook 01) and was **not** used as the GREmLN prior."""))

    c.append(md("""## 10. Supplementary (S1) — held-out seed-TF recovery (NOT a validator)

An internal ranking-consistency diagnostic: hold out a fraction of BTLA seed-TFs and score how each
model recovers them from the remaining seeds. GREmLN's recall@25 saturates at 1.00 here, which is a
seed-clustering artefact; this is **pending a leakage/asymmetry audit** and does **not** enter the
verdict. Left unexecuted by default; enable by setting `RUN_HELDOUT = True`."""))
    c.append(code("""
RUN_HELDOUT = False
if RUN_HELDOUT:
    print("Held-out recovery is supplementary; see the development archive for the full computation.")
else:
    print("Held-out recovery skipped (supplementary; excluded from verdict).")
"""))

    c.append(md("""## 11. Restrained verdict — three-validator decision rule

A model is **superior** only if it wins CRISPRi by ≥ `CRISPRI_MARGIN` on mean functional
target-response **and** has ≥ as many `usable_supportive` TFs **and** carries ≥ corroborating
literature + multi-omics support. Otherwise: equivalence-not-tested, GENIE3-superior, or
complementary. Held-out recovery is excluded.

The primary CRISPRi metric is the **mean target-response fraction over QC-passed perturbations**
(on-target-significance flag), identical to `crispri_model_summary` / Table 2 / Figure 3. The
all-screen mean (including QC failures) is retained only as a labelled sensitivity value."""))
    c.append(code("""
gr_s, g3_s = cr_stats["GREmLN"], cr_stats["GENIE3"]
d_mean = gr_s["mean_frac_8hr"] - g3_s["mean_frac_8hr"]   # QC-passed primary metric

# bootstrap 95% CI on the between-model DIFFERENCE in QC-passed mean target-response fraction
_rng = np.random.default_rng(RANDOM_SEED)
_gr, _ge = _QC_FRAC.get("GREmLN"), _QC_FRAC.get("GENIE3")
if _gr is not None and _ge is not None and len(_gr) and len(_ge):
    _db = [_rng.choice(_gr, len(_gr), replace=True).mean() - _rng.choice(_ge, len(_ge), replace=True).mean()
           for _ in range(5000)]
    diff_ci = [round(float(np.quantile(_db, .025)), 4), round(float(np.quantile(_db, .975)), 4)]
else:
    diff_ci = [None, None]

# corroborating-evidence support among each model's method-SPECIFIC top-25 candidates:
# strong/moderate Paperclip literature + independent orthogonal multi-omics (protein/PTM/coIP).
indep_support = set(integ.loc[integ.get("independent_orthogonal_support", False) == True, "TF"]) \\
    if "independent_orthogonal_support" in integ.columns else set()
corrob_support = strong_mod | indep_support      # union: a TF with both counts once
corrob = {
    "GREmLN": len(gr_only & corrob_support),
    "GENIE3": len(g3_only & corrob_support),
}
# full decision rule: CRISPRi margin win AND >= usable_supportive AND >= corroborating support
gremln_wins = ((d_mean >= CRISPRI_MARGIN)
               and (gr_s["usable_supportive"] >= g3_s["usable_supportive"])
               and (corrob["GREmLN"] >= corrob["GENIE3"]))
genie3_wins = ((-d_mean >= CRISPRI_MARGIN)
               and (g3_s["usable_supportive"] >= gr_s["usable_supportive"])
               and (corrob["GENIE3"] >= corrob["GREmLN"]))
if gremln_wins:
    verdict = "GREmLN superiority"
elif genie3_wins:
    verdict = "GENIE3 superiority"
else:
    verdict = "complementary candidate prioritisation (neither met the decision rule; equivalence not tested)"

prim = overlap["seed_excluded"]
summary = {
    "ranking_primary": "seed_excluded",
    "spearman_rho_seed_excluded": prim["spearman_rank_rho"],
    "shared_top25_seed_excluded": prim["shared_top25"],
    "gremln_only_top25": prim["gremln_only_top25"],
    "genie3_only_top25": prim["genie3_only_top25"],
    "crispri_mean_frac_8hr_qc_passed": {"GREmLN": gr_s["mean_frac_8hr"], "GENIE3": g3_s["mean_frac_8hr"],
                              "difference": round(d_mean, 4), "difference_ci95": diff_ci,
                              "margin": CRISPRI_MARGIN, "population": "QC-passed (primary)"},
    "crispri_mean_frac_8hr_all_screen_sensitivity": {"GREmLN": gr_s["mean_frac_8hr_all_screen"],
                              "GENIE3": g3_s["mean_frac_8hr_all_screen"], "population": "all in-screen (sensitivity)"},
    "crispri_usable_supportive": {"GREmLN": gr_s["usable_supportive"], "GENIE3": g3_s["usable_supportive"]},
    "corroborating_support_method_specific": corrob,
    "decision_rule": "CRISPRi margin AND >= usable_supportive AND >= corroborating (literature+independent multi-omics)",
    "paperclip_coverage_complete": paperclip_coverage_complete,
    "paperclip_literature_support": paperclip_lit,
    "held_out_recovery": "SUPPLEMENTARY — excluded from verdict",
    "verdict": verdict,
}
(REG / "benchmark_summary.json").write_text(json.dumps(summary, indent=2, default=str))
print(json.dumps(summary, indent=2, default=str))
print("\\nVERDICT:", verdict)
"""))

    c.append(md("""## 12. Canonical publication-data export

Freezes the analytical inputs consumed by **notebook 04** into `results/publication_data/`, applying
the two agreed fairness fixes: **(i)** both models score against the *identical common BTLA seed set*
(GREmLN∩GENIE3 reachable), and **(ii)** every common-universe TF is explicitly classified by a
**bona-fide-TF audit** (non-sequence-specific / general-machinery factors are flagged, not silently
dropped; the primary ranking still spans the full common universe). Matched target budgets (native /
top-5 / top-10) and independent-vs-derived multi-omics are exported for the figures/tables. No model
is re-run; this is a re-scoring of frozen GREmLN embeddings and the frozen GENIE3 graph."""))
    c.append(code("""
PUB = bu.repo_root() / "results" / "publication_data"; PUB.mkdir(parents=True, exist_ok=True)

# ---- (a) bona-fide-TF classification over the common universe (annotate, do not drop) ----
NON_BONAFIDE = {
    "CYCS":   ("non_transcriptional", "cytochrome c; electron transport, not a TF"),
    "KIF22":  ("non_transcriptional", "kinesin motor protein, not a TF"),
    "SRP9":   ("non_transcriptional", "signal-recognition-particle subunit, not a TF"),
    "SMAP2":  ("non_transcriptional", "ArfGAP; not a TF"),
    "STAU2":  ("non_transcriptional", "Staufen dsRNA-binding protein; not a TF"),
    "TRAF4":  ("non_transcriptional", "TNFR-associated adaptor; not a TF"),
    "GAR1":   ("non_transcriptional", "H/ACA snoRNP component; not a TF"),
    "HMGN3":  ("chromatin_architectural", "HMGN nucleosome-binding; not sequence-specific"),
    "HIRIP3": ("chromatin_architectural", "histone-chaperone-associated; not sequence-specific"),
    "GTF2A2": ("general_machinery", "general transcription factor IIA; not sequence-specific"),
    "GTF3C2": ("general_machinery", "general transcription factor IIIC; not sequence-specific"),
    "SNAPC4": ("general_machinery", "snRNA-activating complex; not sequence-specific"),
}
def classify_tf(g):
    if g in NON_BONAFIDE:
        cls, reason = NON_BONAFIDE[g]; return cls, reason, False
    return "sequence_specific_or_regulatory_TF", "", True
audit = pd.DataFrame([
    {"TF": g, "in_gremln_universe": g in gremln_tfs, "in_genie3_universe": g in genie3_tfs,
     "is_BTLA_vs_TCR_seed": g in set(seeds), "tf_class": classify_tf(g)[0],
     "is_bona_fide_tf": classify_tf(g)[2], "classification_reason": classify_tf(g)[1]}
    for g in sorted(common)])
audit.to_csv(PUB / "candidate_universe_audit.csv", index=False)
bona_universe = set(audit.loc[audit.is_bona_fide_tf, "TF"])
print(f"common universe={len(common)} | bona-fide={len(bona_universe)} | "
      f"flagged non-bona-fide={len(common) - len(bona_universe)}")

# ---- (b) identical common BTLA seed set for BOTH models ----
graph_nodes = set(edges["regulator"]) | set(edges["target"])
gremln_seeds = set(seeds) & set(genes)
genie3_seeds = set(seeds) & graph_nodes
common_seeds = sorted(gremln_seeds & genie3_seeds)
common_seed_set = set(common_seeds)
seed_univ = pd.DataFrame([
    {"seed_gene": g, "BTLA_vs_TCR_direction": ("BTLA_" + seed_dir[g]) if g in seed_dir else "na",
     "gremln_available": g in set(genes), "genie3_reachable": g in graph_nodes,
     "in_common_seed_set": g in common_seed_set} for g in sorted(seeds)])
seed_univ.to_csv(PUB / "common_seed_universe.csv", index=False)
print(f"panel seeds={len(seeds)} | GREmLN-available={len(gremln_seeds)} | "
      f"GENIE3-reachable={len(genie3_seeds)} | identical common seed set={len(common_seeds)}")
"""))
    c.append(code("""
# ---- (c) re-score BOTH models over the identical common seed set (fair fix #1) ----
gi = {g: i for i, g in enumerate(genes)}
def gremln_neighbour_seeds(tf):
    # common seeds among the TF's top-100 CSLS neighbours, ranked by CSLS similarity (desc)
    if tf not in gi: return []
    row = S_csls[gi[tf]]
    nn = np.argpartition(-row, min(TOPN, len(genes) - 1))[:TOPN]
    nn = nn[np.argsort(-row[nn])]
    return [genes[j] for j in nn if genes[j] in common_seed_set]
def gremln_csls_seed_sum(tf):
    # continuous tie-break: summed CSLS similarity of the common BTLA seeds among the top-100 neighbours
    if tf not in gi: return 0.0
    row = S_csls[gi[tf]]
    nn = np.argpartition(-row, min(TOPN, len(genes) - 1))[:TOPN]
    nn = nn[np.argsort(-row[nn])]
    return float(sum(float(row[j]) for j in nn if genes[j] in common_seed_set))
g3c = edges[edges["target"].isin(common_seed_set)]
g3_sorted = {r: g.sort_values("weight", ascending=False)["target"].tolist() for r, g in g3c.groupby("regulator")}
g3_weight = g3c.groupby("regulator")["weight"].sum()
g3_count = g3c.groupby("regulator")["target"].nunique()   # linked common-BTLA-seed targets (count)
def gremln_score(tf): return len(gremln_neighbour_seeds(tf))
def genie3_score(tf): return float(g3_weight.get(tf, 0.0))
def genie3_n_seed_targets(tf): return int(g3_count.get(tf, 0))   # rankings-neutral display quantity

def build_ranking(exclude_seeds):
    rows = []
    for tf in sorted(common):
        is_seed = tf in set(seeds)
        if exclude_seeds and is_seed: continue
        rows.append({"TF": tf, "gremln_score": gremln_score(tf), "genie3_score": genie3_score(tf),
                     "gremln_csls_seed_sum": round(gremln_csls_seed_sum(tf), 6),
                     "genie3_n_seed_targets": genie3_n_seed_targets(tf),
                     "is_BTLA_vs_TCR_seed": is_seed, "is_bona_fide_tf": tf in bona_universe,
                     "BTLA_vs_TCR_direction": ("BTLA_" + seed_dir[tf]) if tf in seed_dir else "not_seed"})
    d = pd.DataFrame(rows)
    d["gremln_dense_rank"] = rankdata(-d["gremln_score"], method="dense").astype(int)
    d["genie3_dense_rank"] = rankdata(-d["genie3_score"], method="dense").astype(int)
    # dense rank is on the integer seed count; ties are ordered by the continuous CSLS seed-sum (desc)
    return d.sort_values(["gremln_dense_rank", "gremln_csls_seed_sum"], ascending=[True, False]).reset_index(drop=True)

primary = build_ranking(exclude_seeds=True)
context = build_ranking(exclude_seeds=False)
primary.to_csv(PUB / "primary_rankings_common_universe.csv", index=False)
context.to_csv(PUB / "seed_inclusive_rankings.csv", index=False)
assert not primary["is_BTLA_vs_TCR_seed"].any(), "seeds leaked into the primary seed-excluded ranking"

gr25 = list(primary.sort_values(["gremln_dense_rank", "gremln_csls_seed_sum"], ascending=[True, False]).head(25)["TF"])
g325 = list(primary.sort_values("genie3_dense_rank").head(25)["TF"])
union_primary_tfs = sorted(set(gr25) | set(g325))
up = primary[primary.TF.isin(union_primary_tfs)][["TF", "gremln_dense_rank", "genie3_dense_rank",
      "is_BTLA_vs_TCR_seed", "is_bona_fide_tf"]].copy()
up["in_gremln_top25"] = up.TF.isin(gr25); up["in_genie3_top25"] = up.TF.isin(g325)
up["status"] = np.where(up.in_gremln_top25 & up.in_genie3_top25, "shared",
               np.where(up.in_gremln_top25, "GREmLN_specific", "GENIE3_specific"))
up.sort_values(["status", "TF"]).to_csv(PUB / "top25_union_primary.csv", index=False)
rho_primary = spearmanr(primary["gremln_dense_rank"], primary["genie3_dense_rank"])[0]
print(f"primary top-25: GREmLN∩GENIE3 shared={len(set(gr25)&set(g325))} | union={len(union_primary_tfs)} | "
      f"Spearman rho={rho_primary:.3f}")
print("union not in old Paperclip review (if any):",
      sorted(set(union_primary_tfs) - set(pc_join['TF'])))
"""))
    c.append(code("""
# ---- (d) canonical CRISPRi with matched target budgets (native / top5 / top10) ----
def canonical_targets(tf, model):
    if model == "GREmLN": return gremln_neighbour_seeds(tf)
    return [t for t in g3_sorted.get(tf, []) if t in common_seed_set]

def canonical_crispri(union_by_model, de_stats_path, budgets=(None, 5, 10)):
    import anndata as ad, h5py
    all_tfs = sorted({t for v in union_by_model.values() for t in v})
    de = ad.read_h5ad(de_stats_path, backed="r"); de_obs = de.obs.reset_index(drop=True)
    de_gene = de.var["gene_name"].astype(str).values; n_meas = len(de_gene)
    def arm(cond):
        mask = ((de_obs["culture_condition"].astype(str) == cond)
                & (de_obs["target_contrast_gene_name"].astype(str).isin(all_tfs)))
        ri = np.where(mask.values)[0]; order = np.argsort(ri); sidx = ri[order]; inv = np.argsort(order)
        with h5py.File(de_stats_path, "r") as f:
            adjp = f["layers/adj_p_value"][sidx, :][inv]; logfc = f["layers/log_fc"][sidx, :][inv]
        tfn = de_obs.loc[ri, "target_contrast_gene_name"].astype(str).values
        ot = de_obs.loc[ri, "ontarget_significant"].values
        pmap = {}
        for i, t in enumerate(tfn):
            if t not in pmap or (not bool(ot[pmap[t]]) and bool(ot[i])): pmap[t] = i
        return adjp, logfc, ot, pmap
    a8, l8, o8, p8 = arm(STIM_PRIMARY); a48, l48, o48, p48 = arm(STIM_SENS)
    sign = {g: (1 if seed_dir.get(g) == "up" else -1) for g in seeds}
    def de_hits(pmap, adjp, tf):
        if tf not in pmap: return None
        padj = adjp[pmap[tf], :]; hit = np.isfinite(padj) & (padj < DE_PADJ)
        return set(de_gene[hit]), int(hit.sum())
    rows = []
    for model, top in union_by_model.items():
        for tf in top:
            ranked = canonical_targets(tf, model); present = tf in p8
            d8, d48 = de_hits(p8, a8, tf), de_hits(p48, a48, tf)
            r = {"TF": tf, "model": model, "is_BTLA_vs_TCR_seed": tf in set(seeds),
                 "in_screen_8hr": bool(present), "in_screen_48hr": bool(tf in p48),
                 "ontarget_qc_pass": (bool(o8[p8[tf]]) if present else np.nan),
                 "n_predicted_native": len(ranked)}
            for b in budgets:
                lab = "native" if b is None else ("top%d" % b)
                pred = set(ranked if b is None else ranked[:b]); npred = len(pred)
                r["n_predicted_" + lab] = npred
                for al, di in (("8hr", d8), ("48hr", d48)):
                    if di is None or npred == 0:
                        r["n_confirmed_%s_%s" % (lab, al)] = np.nan
                        r["frac_%s_%s" % (lab, al)] = np.nan; r["p_%s_%s" % (lab, al)] = np.nan
                    else:
                        dg, nde = di; conf = pred & dg; k = len(conf)
                        r["n_confirmed_%s_%s" % (lab, al)] = k; r["frac_%s_%s" % (lab, al)] = k / npred
                        r["p_%s_%s" % (lab, al)] = float(hypergeom.sf(k - 1, n_meas, nde, npred)) if k > 0 else np.nan
                        if al == "8hr": r["confirmed_targets_" + lab] = "|".join(sorted(conf))
            na = no = 0
            if present and d8:
                lfc = dict(zip(de_gene, l8[p8[tf], :]))
                for g in (set(ranked) & d8[0]):
                    lf = lfc.get(g, np.nan)
                    if not np.isfinite(lf) or lf == 0 or g not in sign: continue
                    na += int(np.sign(lf) == np.sign(sign[g])); no += int(np.sign(lf) != np.sign(sign[g]))
            net = (na - no) / (na + no) if (na + no) else np.nan
            r["n_kd_agrees_btla"] = na; r["n_kd_opposes_btla"] = no; r["net_signed_agreement"] = net
            r["response_direction"] = ("no_response" if not (na + no) else
                "BTLA_concordant" if net >= 0.5 else "anti_concordant" if net <= -0.5 else "mixed")
            rows.append(r)
    df = pd.DataFrame(rows)
    # keep a raw in-screen fraction (incl. QC-fail) for the all-screen sensitivity arm, then treat
    # failed-on-target-QC perturbations as NOT SCORED (NaN) so they are never counted as a zero response.
    df["frac_native_8hr_all_screen"] = df["frac_native_8hr"]
    _null = [x for x in df.columns if x.startswith(("n_confirmed_", "frac_", "p_", "confirmed_targets_"))
             and x != "frac_native_8hr_all_screen"]
    df.loc[df["ontarget_qc_pass"] == False, _null] = np.nan
    df["adj_p_native_8hr"] = np.nan
    for m in df.model.unique():
        idx = df.model == m; p = df.loc[idx, "p_native_8hr"].values; ok = np.isfinite(p)
        if ok.sum():
            adj = np.full(len(p), np.nan); adj[ok] = multipletests(p[ok], method="fdr_bh")[1]
            df.loc[idx, "adj_p_native_8hr"] = adj
    def status(r):
        if not r["in_screen_8hr"]: return "not_in_screen"
        if r["ontarget_qc_pass"] == False: return "failed_ontarget_qc"
        n, p, net = r["n_confirmed_native_8hr"], r["p_native_8hr"], r["net_signed_agreement"]
        if pd.notna(n) and n >= 3 and pd.notna(net) and net <= -0.5: return "usable_contradictory"
        if pd.notna(n) and n > 0 and pd.notna(p) and p < 0.05: return "usable_supportive"
        return "usable_unsupported"
    df["validation_status"] = df.apply(status, axis=1)
    return df

crispri_canon = canonical_crispri({"GREmLN": gr25, "GENIE3": g325}, bu.artifact("de_stats"))
crispri_canon.to_csv(PUB / "crispri_all_screen_sensitivity.csv", index=False)
qc = crispri_canon[(crispri_canon.in_screen_8hr) & (crispri_canon.ontarget_qc_pass == True)].copy()
qc.to_csv(PUB / "crispri_primary_qc_passed.csv", index=False)
budget_cols = ["TF", "model", "in_screen_8hr", "ontarget_qc_pass"]
crispri_canon[budget_cols + [c for c in crispri_canon.columns if "top5" in c]].to_csv(PUB / "crispri_matched_budget_top5.csv", index=False)
crispri_canon[budget_cols + [c for c in crispri_canon.columns if "top10" in c]].to_csv(PUB / "crispri_matched_budget_top10.csv", index=False)

def model_summary(m):
    sub = crispri_canon[crispri_canon.model == m]
    q = sub[(sub.in_screen_8hr) & (sub.ontarget_qc_pass == True)]
    fr = q["frac_native_8hr"].astype(float).dropna()
    rng = np.random.default_rng(RANDOM_SEED)
    boot = [rng.choice(fr.values, len(fr), replace=True).mean() for _ in range(5000)] if len(fr) else [np.nan]
    d = {"model": m, "present_in_screen": int(sub.in_screen_8hr.sum()),
         "qc_passed": int((sub.ontarget_qc_pass == True).sum()), "n_in_primary": int(len(q)),
         "mean_frac_native_8hr": round(float(fr.mean()), 4) if len(fr) else np.nan,
         "median_frac_native_8hr": round(float(fr.median()), 4) if len(fr) else np.nan,
         "ci_lo": round(float(np.nanquantile(boot, .025)), 4), "ci_hi": round(float(np.nanquantile(boot, .975)), 4),
         "mean_frac_top5_8hr": round(float(q["frac_top5_8hr"].astype(float).mean()), 4),
         "mean_frac_top10_8hr": round(float(q["frac_top10_8hr"].astype(float).mean()), 4),
         "mean_frac_all_screen_native_8hr": round(float(sub["frac_native_8hr_all_screen"].astype(float).mean()), 4)}
    for s in ["usable_supportive", "usable_unsupported", "usable_contradictory", "failed_ontarget_qc", "not_in_screen"]:
        d[s] = int((sub.validation_status == s).sum())
    for rd in ["BTLA_concordant", "anti_concordant", "mixed", "no_response"]:
        d["dir_" + rd] = int((q.response_direction == rd).sum())
    return d
crispri_model = pd.DataFrame([model_summary("GREmLN"), model_summary("GENIE3")])
crispri_model.to_csv(PUB / "crispri_model_summary.csv", index=False)
print(crispri_model[["model", "present_in_screen", "qc_passed", "mean_frac_native_8hr", "ci_lo", "ci_hi",
                     "mean_frac_top5_8hr", "mean_frac_top10_8hr", "usable_supportive"]].to_string(index=False))
"""))
    c.append(code("""
# ---- (e) evidence joins (annotation only): Paperclip + independent-vs-derived multi-omics ----
INDEP = ["protein_support", "phosphosite_support", "coip_support"]
DERIV = ["transcript_support", "tf_activity_support", "kinase_activity_support",
         "bionic_support", "early_synapse_trafficking_support"]
uni = pd.DataFrame({"TF": union_primary_tfs}).merge(
    up[["TF", "status", "gremln_dense_rank", "genie3_dense_rank", "is_BTLA_vs_TCR_seed", "is_bona_fide_tf"]],
    on="TF", how="left")
pc_full = pd.read_csv(bu.repo_root() / "results" / "paperclip" / "paperclip_union_top25_review.csv")
pc_cols = [c for c in ["TF", "paperclip_evidence_tier", "paperclip_primary_phenotype", "paperclip_direction",
           "tcr_checkpoint_evidence", "activation_exhaustion_evidence", "btla_specific_evidence",
           "paperclip_key_source", "paperclip_short_rationale", "paperclip_reviewed"] if c in pc_full.columns]
paperclip_candidate = uni.merge(pc_full[pc_cols], on="TF", how="left")
paperclip_candidate["paperclip_reviewed"] = paperclip_candidate["paperclip_reviewed"].fillna(False).astype(bool)
paperclip_candidate.to_csv(PUB / "paperclip_candidate_evidence.csv", index=False)

mo_path = bu.artifact("multiomics_summary")
mo_out = uni[["TF", "status", "is_BTLA_vs_TCR_seed"]].copy()
for col in INDEP + DERIV:
    mo_out[col] = np.nan
if mo_path.exists():
    mo = pd.read_csv(mo_path).rename(columns={"gene": "TF"})
    keep = ["TF"] + [c for c in INDEP + DERIV if c in mo.columns]
    mo_out = uni[["TF", "status", "is_BTLA_vs_TCR_seed"]].merge(mo[keep], on="TF", how="left")
# "not_significant"/"not_measured" mean tested-without-support and must NOT count as evidence
# (fixes e.g. RBPJ protein_support="not_significant" being mis-scored as independent support).
def _has(v): return str(v).lower() not in ("", "nan", "none", "no", "false", "0", "0.0",
                                           "not_significant", "not_measured", "hypothesis_only")
mo_out["independent_orthogonal_support"] = mo_out[[c for c in INDEP if c in mo_out.columns]].apply(
    lambda r: any(_has(x) for x in r), axis=1)
mo_out["derived_contextual_support"] = mo_out[[c for c in DERIV if c in mo_out.columns]].apply(
    lambda r: any(_has(x) for x in r), axis=1)
mo_out.to_csv(PUB / "multiomics_candidate_evidence.csv", index=False)

# ---- long-form per-observation evidence (full provenance) for the candidate union ----
# Tables 4A/4B and Supplementary S2 are generated from THIS, never from hand-typed labels.
lf_path = bu.artifact("multiomics_long")
LF_COLS = ["gene", "evidence_layer", "contrast", "timepoint", "condition", "direction",
           "effect_size", "adjusted_p", "tf_activity_score", "kinase_activity_score",
           "bionic_cluster", "phosphosite", "ip_condition", "source_file",
           "direct_or_inferred", "independent_or_derived", "support_level", "notes"]
if lf_path.exists():
    lf = pd.read_csv(lf_path)
    lf = lf[lf["gene"].isin(union_primary_tfs)].copy()
    for c in LF_COLS:
        if c not in lf.columns:
            lf[c] = np.nan
    lf = lf[LF_COLS].sort_values(["gene", "evidence_layer", "contrast", "timepoint"]).reset_index(drop=True)
else:
    lf = pd.DataFrame(columns=LF_COLS)
    print("WARNING: multiomics_long artifact missing:", lf_path)
lf.to_csv(PUB / "multiomics_long_evidence.csv", index=False)
print("long-form evidence rows for union candidates:", len(lf),
      "| candidates covered:", lf["gene"].nunique())

# integrated evidence map (grouped, NOT collapsed into one score)
cr8 = crispri_canon.set_index(["TF", "model"])
integ_rows = []
for _, r in uni.iterrows():
    tf = r["TF"]; model = "GREmLN" if r["status"] in ("shared", "GREmLN_specific") else "GENIE3"
    key = (tf, model) if (tf, model) in cr8.index else ((tf, "GREmLN") if (tf, "GREmLN") in cr8.index else (tf, "GENIE3") if (tf, "GENIE3") in cr8.index else None)
    cr = cr8.loc[key] if key in cr8.index else None
    pc = paperclip_candidate.set_index("TF").loc[tf]; mo = mo_out.set_index("TF").loc[tf]
    integ_rows.append({
        "TF": tf, "status": r["status"], "is_BTLA_vs_TCR_seed": r["is_BTLA_vs_TCR_seed"],
        "is_bona_fide_tf": r["is_bona_fide_tf"], "gremln_rank": r["gremln_dense_rank"], "genie3_rank": r["genie3_dense_rank"],
        "crispri_status": (cr["validation_status"] if cr is not None else "not_in_screen"),
        "ontarget_qc_pass": (cr["ontarget_qc_pass"] if cr is not None else np.nan),
        "frac_native_8hr": (cr["frac_native_8hr"] if cr is not None else np.nan),
        "response_direction": (cr["response_direction"] if cr is not None else "na"),
        "paperclip_evidence_tier": pc.get("paperclip_evidence_tier", np.nan),
        "protein_support": mo.get("protein_support", np.nan), "phosphosite_support": mo.get("phosphosite_support", np.nan),
        "coip_support": mo.get("coip_support", np.nan), "transcript_support": mo.get("transcript_support", np.nan),
        "tf_activity_support": mo.get("tf_activity_support", np.nan), "kinase_activity_support": mo.get("kinase_activity_support", np.nan),
        "bionic_support": mo.get("bionic_support", np.nan),
        "independent_orthogonal_support": mo.get("independent_orthogonal_support", False),
        "derived_contextual_support": mo.get("derived_contextual_support", False)})
candidate_integrated = pd.DataFrame(integ_rows)
candidate_integrated.to_csv(PUB / "candidate_integrated_evidence.csv", index=False)
print("evidence rows:", len(candidate_integrated), "| paperclip reviewed:",
      int(paperclip_candidate.paperclip_reviewed.sum()), "/", len(paperclip_candidate))
"""))
    c.append(code("""
# ---- (f) model/benchmark specification + manifest ----
n_edges = len(edges)
spec = pd.DataFrame([
    ("biological_task", "BTLA+TCR vs TCR 4h TF prioritisation (249-DEG panel)", "same for both", "identical target task"),
    ("cells", "CD4 NTC (CZI Perturb-seq)", "same cells", "same cells"),
    ("expression_input", "donor-ComBat log1p-CPM (masked)", "pre-ComBat log1p-CPM (same 50k cells)", "different numerical inputs"),
    ("preprocessing", "HVG-4000 -> CPM/log1p -> ComBat(donor)", "tokenizer rank/zero structure", "GREmLN needs raw-like structure"),
    ("graph_prior", "none (co-expression forest)", "masked GENIE3 50k-edge graph", "GREmLN is handed GENIE3's graph"),
    ("graph_size", str(n_edges) + " edges", "same graph as prior", "shared topology"),
    ("model_checkpoint", "GENIE3 (ExtraTrees)", "GREmLN model.ckpt (zero-shot)", "no fine-tuning"),
    ("gene_universe", str(len(set(edges['regulator'])|set(edges['target']))) + " graph genes", str(len(genes)) + " embedded genes", "coverage differs"),
    ("common_seed_universe", str(len(common_seeds)) + " identical common seeds", str(len(common_seeds)), "fair fix #1"),
    ("common_bona_fide_TF_universe", str(len(bona_universe)) + " of " + str(len(common)) + " common TFs", "same", "bona-fide audit"),
    ("scoring_method", "summed outgoing TF->seed edge weight", "CSLS top-100 seed-neighbour count (k=10)", "unsigned, non-comparable raw scores -> ranks"),
    ("seed_handling", "seed-excluded primary; seed-inclusive context", "same", "identical rule"),
    ("predicted_target_definition", "graph edges into common seeds", "top-100 CSLS neighbours in common seeds", "model-specific target sets"),
    ("crispri_validation_arm", "Stim8hr primary; Stim48hr sensitivity", "same", "orthogonal functional readout"),
    ("primary_fairness_limitation", "supplies GENIE3 graph to GREmLN; different expression inputs", "same", "measures embedding value given the prior, not two independent methods"),
], columns=["characteristic", "GENIE3", "GREmLN", "implication_for_comparison"])
spec.to_csv(PUB / "model_benchmark_specification.csv", index=False)

files = ["model_benchmark_specification", "common_seed_universe", "candidate_universe_audit",
         "primary_rankings_common_universe", "seed_inclusive_rankings", "top25_union_primary",
         "crispri_primary_qc_passed", "crispri_all_screen_sensitivity", "crispri_matched_budget_top5",
         "crispri_matched_budget_top10", "crispri_model_summary", "paperclip_candidate_evidence",
         "multiomics_candidate_evidence", "multiomics_long_evidence", "candidate_integrated_evidence"]
man = []
for name in files:
    p = PUB / (name + ".csv"); d = pd.read_csv(p)
    man.append({"file": name + ".csv", "rows": len(d), "cols": d.shape[1], "md5": bu.md5(p)})
manifest = pd.DataFrame(man)
manifest.to_csv(PUB / "publication_data_manifest.csv", index=False)
print(manifest.to_string(index=False))
print("\\ncanonical publication-data package written to", PUB)
"""))

    _write(_nb(c), NB / "03_benchmark_btla_tf_prioritisation.ipynb")


# ============================================================ NOTEBOOK 01
def build_nb1():
    c = []
    c.append(md("""
# 01 — Build the masked CD4 GENIE3 graph (regulatory prior for GREmLN)

Constructs the **raw masked CD4 GENIE3** co-expression graph used both as the GENIE3 comparator and
as GREmLN's regulatory prior (notebook 02). Steps: CD4 **non-targeting-control (NTC)** cell
selection → donor/batch audit → CPM/log1p normalisation → **ComBat** donor masking → TF regulator
universe → **GENIE3** (tree ensembles) → top-50,000-edge graph → model registry, plus a **SCENIC
top-50k sensitivity** (motif support only; **not** the GREmLN prior).

The GENIE3 forest fit is hours-long, so it is behind `REGENERATE`. By default the notebook
**reuses** the persisted `grn_edges.tsv` and audits it. Parameters come from
`config/benchmark_config.yaml`; see `data/README.md` to obtain the CD4 Perturb-seq NTC cells.
"""))
    c.append(code("""
import sys, json
from pathlib import Path
import numpy as np, pandas as pd
_here = Path.cwd()
_root = next((p for p in [_here, *_here.parents] if (p / "scripts" / "bench_utils.py").exists()), _here.parent)
sys.path.insert(0, str(_root / "scripts"))
import bench_utils as bu
import yaml

CFG = yaml.safe_load((bu.repo_root() / "config" / "benchmark_config.yaml").read_text())
REGENERATE = False   # True re-runs GENIE3 (hours; needs the CD4 NTC expression matrix)
REG = bu.repo_root() / "results" / "model_registry"; REG.mkdir(parents=True, exist_ok=True)
print("GENIE3 params:", CFG["genie3"]); print("DATA_ROOT:", bu.data_root())
"""))
    c.append(md("""## 1. Source data, NTC selection, donor/batch audit, normalisation, ComBat masking

The GENIE3 input is CD4 **NTC** cells from the CZI genome-scale T-cell Perturb-seq resource:
select NTC guides, keep CD4, take the top-4,000 highly variable genes, normalise to CPM(1e6)→log1p,
then apply **ComBat on `donor_id`** (this donor-masking defines the *masked* graph). The exact
preprocessing lives in the development archive's prep script; parameters are pinned in the config.
The GREmLN prior in notebook 02 deliberately uses the *pre-ComBat* log1p-CPM of the same cells."""))
    c.append(code("""
# Documented preprocessing (guarded). In reuse mode we do not need the raw matrix.
if REGENERATE:
    import anndata as ad, scanpy as sc
    adata = ad.read_h5ad(bu.artifact("cd4_ntc_counts"))          # provide via BTLA_BENCH_CD4_NTC_COUNTS
    adata = adata[adata.obs["guide_target"].astype(str).eq("non-targeting")].copy()
    sc.pp.highly_variable_genes(adata, n_top_genes=CFG["genie3"]["n_hvg"], flavor="seurat_v3")
    adata = adata[:, adata.var["highly_variable"]].copy()
    sc.pp.normalize_total(adata, target_sum=1e6); sc.pp.log1p(adata)
    sc.pp.combat(adata, key="donor_id")                          # donor masking -> masked graph
    donor_audit = adata.obs["donor_id"].value_counts()
    print("cells x genes:", adata.shape); print(donor_audit)
else:
    print("REGENERATE=False -> reuse persisted masked graph; preprocessing shown above for provenance.")
"""))
    c.append(md("""## 2. TF regulator universe + GENIE3 (tree-ensemble feature importances)

GENIE3 fits, for each target gene, a tree ensemble predicting it from the TF regulators; edge weight
= regulator feature importance. Core algorithm inlined below; the full VIM over ~4k genes is the
hours-long step."""))
    c.append(code('''
def genie3_single(expr, target_idx, reg_idx, ntrees=1000, K="sqrt"):
    # Predict ONE target gene from the regulator set; return per-regulator importances.
    from sklearn.ensemble import ExtraTreesRegressor
    y = np.nan_to_num(np.asarray(expr[:, target_idx], float)); s = y.std()
    if s > 0: y = y / s
    preds = [i for i in reg_idx if i != target_idx]   # regulators are the predictors
    Xr = np.nan_to_num(np.asarray(expr[:, preds], float))
    est = ExtraTreesRegressor(n_estimators=ntrees, max_features=K, n_jobs=1).fit(Xr, y)
    imp = np.array([e.tree_.compute_feature_importances(normalize=False) for e in est.estimators_])
    vi = np.zeros(expr.shape[1]); vi[preds] = imp.sum(0) / len(est); return vi

def run_genie3(expr, gene_names, regulators, n_edges_keep=50000):
    # Standard GENIE3: every gene is a candidate TARGET predicted from the regulators;
    # edge = regulator -> target, weight = regulator's feature importance for that target.
    ng = expr.shape[1]; ridx = [i for i, g in enumerate(gene_names) if g in regulators]
    VIM = np.zeros((ng, ng))                          # VIM[target, regulator]
    for t in range(ng):                              # loop over ALL genes as targets
        VIM[t, :] = genie3_single(expr, t, ridx)
    links = [(gene_names[r], gene_names[t], float(VIM[t, r]))   # regulator r -> target t
             for t in range(ng) for r in ridx if r != t and VIM[t, r] > 0]
    links.sort(key=lambda x: x[2], reverse=True)
    return pd.DataFrame(links[:n_edges_keep], columns=["regulator", "target", "weight"])

if REGENERATE:
    tfs = bu.load_tfs(); gnames = list(adata.var_names.astype(str))
    regs = [g for g in gnames if g in tfs]
    edges = run_genie3(np.asarray(adata.X, float), gnames, regs, CFG["genie3"]["n_edges_persisted"])
    edges.to_csv(bu.artifact("genie3_edges"), sep="\\t", index=False)
else:
    print("Skipping GENIE3 fit (reuse mode).")
'''))
    c.append(md("## 3. Load / audit the persisted masked graph + model registry"))
    c.append(code("""
edges = bu.load_edges()
regs, tgts = set(edges["regulator"]), set(edges["target"])
tfs = bu.load_tfs()
stats = {"n_edges": len(edges), "n_regulators": len(regs), "n_regulators_in_tf_list": len(regs & tfs),
         "n_targets": len(tgts), "n_nodes": len(regs | tgts),
         "weight_min": float(edges["weight"].min()), "weight_max": float(edges["weight"].max())}
print(json.dumps(stats, indent=2))
registry = pd.DataFrame([
    ("model_name", "GENIE3_CD4_masked_raw"),
    ("graph_path", bu.redact(bu.artifact("genie3_edges"))),
    ("graph_md5", bu.md5(bu.artifact("genie3_edges"))),
    ("n_edges", stats["n_edges"]), ("n_regulators", stats["n_regulators"]),
    ("n_targets", stats["n_targets"]),
    ("n_hvg", CFG["genie3"]["n_hvg"]), ("normalisation", CFG["genie3"]["normalisation"]),
    ("batch_correction", CFG["genie3"]["batch_correction"]),
    ("tf_list", CFG["genie3"]["tf_list"]),
    ("tf_list_md5", bu.md5(bu.repo_root() / "resources" / "human_tfs_pySCENIC.txt")),
    ("tree_method", "ExtraTrees/RandomForest, ntrees=1000, K=sqrt"),
], columns=["field", "value"])
registry.to_csv(REG / "genie3_model_registry.csv", index=False)
registry
"""))
    c.append(md("""## 4. SCENIC top-50k sensitivity (motif support only — NOT the GREmLN prior)

> SCENIC pruning of the persisted top-50k GENIE3 graph is retained as an exploratory motif-support
> sensitivity analysis and is **not** used as the GREmLN prior. Enable with `REGENERATE` + pySCENIC
> databases; the pruned graph is reported separately and never fed to GREmLN."""))
    c.append(code("""
if REGENERATE:
    print("Run pySCENIC prune_targets on grn_edges.tsv with cisTarget motif rankings (see archive).")
else:
    print("SCENIC sensitivity skipped; documented as a separate motif-support analysis only.")
"""))
    _write(_nb(c), NB / "01_build_genie3_cd4_masked.ipynb")


# ============================================================ NOTEBOOK 02
def build_nb2():
    c = []
    c.append(md("""
# 02 — GREmLN gene embeddings from the masked GENIE3 prior

Runs the **GREmLN** foundation model (frozen checkpoint, zero-shot) to produce per-gene embeddings,
using the masked CD4 GENIE3 graph (notebook 01) as its **regulatory prior**. Key design point: the
expression fed to GREmLN is the **pre-ComBat log1p-CPM** of the same 50k cells, because GREmLN's
tokenizer needs the rank/zero structure of near-count data (ComBat output has negatives / few true
zeros). GPU inference is behind `REGENERATE`; by default we **reuse** the saved embeddings and audit
coverage. Obtain the checkpoint per `data/README.md`.
"""))
    c.append(code("""
import sys, json, platform
from pathlib import Path
import numpy as np, pandas as pd
_here = Path.cwd()
_root = next((p for p in [_here, *_here.parents] if (p / "scripts" / "bench_utils.py").exists()), _here.parent)
sys.path.insert(0, str(_root / "scripts"))
import bench_utils as bu
import yaml
CFG = yaml.safe_load((bu.repo_root() / "config" / "benchmark_config.yaml").read_text())
REGENERATE = False       # True runs GREmLN inference (GPU). False reuses saved embeddings.
REG = bu.repo_root() / "results" / "model_registry"; REG.mkdir(parents=True, exist_ok=True)
print("GREmLN cfg:", {k: CFG["gremln"][k] for k in ("embedding_dim", "expression_input", "prior_graph")})
"""))
    c.append(md("""## 1. Provenance: checkpoint, submodule commit, expression input, graph prior"""))
    c.append(code("""
def pkg(n):
    try:
        from importlib.metadata import version; return version(n)
    except Exception: return "unknown"
def git_commit(path):
    import subprocess
    try:
        return subprocess.check_output(["git", "rev-parse", "HEAD"], cwd=str(path),
                                       text=True, stderr=subprocess.DEVNULL).strip()
    except Exception: return "unknown"
gremln_commit = git_commit(bu.repo_root() / "third_party" / "GREmLN")
print("GREmLN submodule commit:", gremln_commit)
print("checkpoint md5 (config):", CFG["gremln"]["checkpoint_md5"])
"""))
    c.append(md("""## 2. GREmLN inference (GPU, guarded): vocab → RegulatoryNetwork(prior) → tokenizer → embeddings

The real inference call is inlined below. It loads the frozen `GDTransformer`, builds the
`GraphTokenizer` from the GREmLN gene vocabulary and the masked GENIE3 graph as `RegulatoryNetwork`,
and runs a single gene-embedding forward pass (`get_gene_embeddings`), mapping Ensembl→HUGO."""))
    c.append(code('''
if REGENERATE:
    import anndata as ad, torch
    sys.path.insert(0, str(bu.repo_root() / "third_party" / "GREmLN"))
    from scGraphLLM import GeneVocab, GraphTokenizer, InferenceDataset, RegulatoryNetwork
    from scGraphLLM.config import graph_kernel_attn_3L_4096
    from scGraphLLM.inference import get_gene_embeddings
    from scGraphLLM.models import GDTransformer
    adata = ad.read_h5ad(bu.artifact("gremln_inputs") / "expr.h5ad")     # pre-ComBat log1p-CPM, Ensembl var
    vocab = GeneVocab.from_csv(bu.repo_root() / "third_party/GREmLN/scGraphLLM/resources/gene_vocab.csv",
                               gene_col="gene_name", node_col="idx")
    network = RegulatoryNetwork.from_csv(bu.artifact("gremln_inputs") / "graph.tsv", sep="\\t",
                                         reg_name="regulator.values", tar_name="target.values",
                                         wt_name="mi.values", lik_name="log.p.values")
    model = GDTransformer.load_from_checkpoint(str(bu.artifact("gremln_checkpoint")),
                                               config=graph_kernel_attn_3L_4096).eval().cuda()
    ds = InferenceDataset(expression=adata.to_df(), tokenizer=GraphTokenizer(vocab=vocab, network=network))
    x_gene = get_gene_embeddings(ds, model, vocab, batch_size=256)       # per-gene mean embeddings
    emb = ad.AnnData(x_gene.values, obs=adata.var.loc[x_gene.index].copy())
    emb.obs["hugo"] = emb.obs.get("hugo_symbol", pd.Series(emb.obs_names)).astype(str)
    emb.write_h5ad(bu.artifact("gremln_emb"))
else:
    print("REGENERATE=False -> reuse saved GREmLN embeddings (inference code above for provenance).")
'''))
    c.append(md("## 3. Load embeddings + coverage audit + sanitized model registry"))
    c.append(code("""
genes, X = bu.load_gremln_embeddings()
tfs = bu.load_tfs(); edges = bu.load_edges(); seeds, seed_dir, _ = bu.load_seeds()
gremln_tfs = set(genes) & tfs; genie3_tfs = set(edges["regulator"]) & tfs
common = gremln_tfs & genie3_tfs
audit = pd.DataFrame([
    ("genes_embedded_by_gremln", len(genes)), ("embedding_dim", X.shape[1]),
    ("tfs_embedded_by_gremln", len(gremln_tfs)), ("tfs_in_genie3_graph", len(genie3_tfs)),
    ("common_tf_universe", len(common)), ("btla_seed_panel_size", len(seeds)),
    ("seeds_available_to_gremln", len(set(seeds) & set(genes))),
    ("seed_tfs_in_common_universe", len(set(seeds) & common)),
], columns=["metric", "value"])
print(audit.to_string(index=False))
registry = pd.DataFrame([
    ("model_name", "GREmLN_CD4_masked_raw_prior"),
    ("gremln_checkpoint", "model.ckpt"), ("gremln_checkpoint_md5", CFG["gremln"]["checkpoint_md5"]),
    ("gremln_submodule_commit", gremln_commit),
    ("python_version", platform.python_version()),
    ("numpy_version", pkg("numpy")), ("pandas_version", pkg("pandas")),
    ("anndata_version", pkg("anndata")), ("torch_version", pkg("torch")),
    ("expression_input", CFG["gremln"]["expression_input"]),
    ("prior_graph", CFG["gremln"]["prior_graph"]),
    ("prior_graph_md5", bu.md5(bu.artifact("genie3_edges"))),
    ("embedding_h5ad_md5", bu.md5(bu.artifact("gremln_emb"))),
    ("n_genes_embedded", len(genes)), ("embedding_dim", X.shape[1]),
    ("csls_params", CFG["scoring"]["gremln"]), ("random_seed", CFG["seed"]),
    ("device", "cuda"), ("inference_params", CFG["gremln"]["inference"]),
], columns=["field", "value"])
registry.to_csv(REG / "gremln_model_registry.csv", index=False)
audit.to_csv(REG / "gremln_gene_universe_audit.csv", index=False)
registry
"""))
    _write(_nb(c), NB / "02_build_gremln_from_masked_genie3_prior.ipynb")


def _write(nb, path):
    nbf.write(nb, str(path))
    print("wrote", path)


# ------------------------------------------------------------ NB04 figures
def _build_nb4_figures(c):
    c.append(md("## 2. Figure 1 — study design and fairness framework"))
    c.append(code("""
n_panel = len(seedu); n_gre = int(seedu.gremln_available.sum()); n_gen = int(seedu.genie3_reachable.sum())
n_cs = int(seedu.in_common_seed_set.sum()); n_univ = len(au); n_bona = int(au.is_bona_fide_tf.sum())
n_prim = len(primary); n_ctx = len(D["seed_inclusive_rankings"]); union = D["top25_union_primary"]; n_union = len(union)
cr = D["crispri_all_screen_sensitivity"]
n_screen = int(cr[cr.in_screen_8hr == True].TF.nunique()); n_qc = int(D["crispri_primary_qc_passed"].TF.nunique())
fig1_src = pd.DataFrame([
    ("CD4 Perturb-seq NTC cells", ""), ("Masked CD4 GENIE3 graph", "50,000 edges"),
    ("GREmLN (masked GENIE3 prior; pre-ComBat log1p-CPM input)", ""),
    ("BTLA+TCR vs TCR panel (4 h)", "%d DEGs" % n_panel),
    ("GREmLN-available seeds", str(n_gre)), ("GENIE3-reachable seeds", str(n_gen)),
    ("Identical common seed set", str(n_cs)), ("Initial common regulator universe", str(n_univ)),
    ("Bona-fide common TF universe", str(n_bona)), ("Seed-excluded primary universe", str(n_prim)),
    ("Seed-inclusive contextual universe", str(n_ctx)), ("Top-25 union (both models)", str(n_union)),
    ("CRISPRi in screen", str(n_screen)), ("CRISPRi on-target QC passed", str(n_qc)),
    ("Paperclip + multi-omics corroboration", "annotation only"),
], columns=["step", "count"])
print("source data for Figure 1:"); display(fig1_src)
"""))
    c.append(code("""
def _box(ax, x, y, w, h, text, fc):
    ax.add_patch(FancyBboxPatch((x - w / 2, y - h / 2), w, h, boxstyle="round,pad=0.02,rounding_size=0.04",
                                fc=fc, ec="#333333", lw=1.2))
    ax.text(x, y, text, ha="center", va="center", fontsize=8.4)
def _arw(ax, p1, p2):
    ax.add_patch(FancyArrowPatch(p1, p2, arrowstyle="-|>", mutation_scale=11, color="#666666", lw=1.0))

INP, UNI, VAL = "#DCE9F5", "#E8F5E9", "#FDEBD0"
fig, ax = plt.subplots(figsize=(9.6, 12)); ax.set_xlim(0, 10); ax.set_ylim(0, 15); ax.axis("off")
_box(ax, 5, 14.3, 6.4, 0.8, "CD4 Perturb-seq NTC cells", INP)
_box(ax, 5, 13.1, 6.4, 0.8, "Masked CD4 GENIE3 graph  —  50,000 edges", INP)
_box(ax, 5, 11.85, 7.6, 1.0, "GREmLN (frozen): masked GENIE3 graph as prior;\\npre-ComBat log1p-CPM expression input", INP)
_box(ax, 5, 10.6, 6.4, 0.8, "BTLA+TCR vs TCR panel (4 h)  —  %d DEGs" % n_panel, INP)
_box(ax, 2.7, 9.35, 4.2, 0.8, "GREmLN-available seeds  —  %d" % n_gre, UNI)
_box(ax, 7.3, 9.35, 4.2, 0.8, "GENIE3-reachable seeds  —  %d" % n_gen, UNI)
_box(ax, 5, 8.1, 5.2, 0.8, "Identical common seed set  —  %d" % n_cs, UNI)
_box(ax, 5, 6.9, 6.8, 0.8, "Initial common regulator universe  —  %d" % n_univ, UNI)
_box(ax, 5, 5.7, 6.8, 0.8, "Bona-fide common TF universe  —  %d" % n_bona, UNI)
_box(ax, 2.7, 4.5, 4.2, 0.8, "Seed-excluded PRIMARY  —  %d" % n_prim, UNI)
_box(ax, 7.3, 4.5, 4.2, 0.8, "Seed-inclusive context  —  %d" % n_ctx, UNI)
_box(ax, 5, 3.3, 5.4, 0.8, "Top-25 union (both models)  —  %d" % n_union, UNI)
_box(ax, 5, 2.05, 7.4, 1.0, "CD4 CRISPRi functional validation\\n(in screen %d; on-target QC %d)" % (n_screen, n_qc), VAL)
_box(ax, 5, 0.75, 7.4, 0.8, "Paperclip + multi-omics corroboration (annotation only)", VAL)
for a, b in [(14.3, 13.1), (13.1, 11.85), (11.85, 10.6)]: _arw(ax, (5, a - 0.42), (5, b + 0.5))
_arw(ax, (5, 10.2), (2.7, 9.75)); _arw(ax, (5, 10.2), (7.3, 9.75))
_arw(ax, (2.7, 8.95), (5, 8.5)); _arw(ax, (7.3, 8.95), (5, 8.5))
for a, b in [(8.1, 6.9), (6.9, 5.7)]: _arw(ax, (5, a - 0.42), (5, b + 0.42))
_arw(ax, (5, 5.3), (2.7, 4.9)); _arw(ax, (5, 5.3), (7.3, 4.9))
_arw(ax, (2.7, 4.1), (5, 3.72))
for a, b in [(3.3, 2.05), (2.05, 0.75)]: _arw(ax, (5, a - 0.42), (5, b + 0.5))
ax.set_title("Study design and fairness framework for the GREmLN\\u2013GENIE3 benchmark", fontsize=12, pad=6)
ax.legend(handles=[Line2D([0], [0], marker="s", color="w", markerfacecolor=INP, markersize=12, label="inputs"),
                   Line2D([0], [0], marker="s", color="w", markerfacecolor=UNI, markersize=12, label="analysis universe"),
                   Line2D([0], [0], marker="s", color="w", markerfacecolor=VAL, markersize=12, label="validation / evidence")],
          loc="upper right", fontsize=8)
save_fig(fig, "fig1_study_design", fig1_src,
    "Figure 1. Study design and fairness framework. Boxes are analysis stages coloured by role "
    "(blue inputs, green analysis universe, orange validation/evidence); numbers are exact counts of "
    "cells, seeds or TFs entering each stage. Both models are restricted to the identical common seed "
    "set and the same common TF universe; CD4 CRISPRi is the functional validator; Paperclip and "
    "multi-omics provide corroboration only, never predictive input.")
plt.show()
"""))

    c.append(md("## 3. Figure 2 — primary model-ranking comparison"))
    c.append(code("""
prim_b = primary[primary.is_bona_fide_tf].copy()
union = D["top25_union_primary"].copy(); status_map = dict(zip(union.TF, union.status))
fig2_src = union[["TF", "status", "gremln_dense_rank", "genie3_dense_rank", "is_BTLA_vs_TCR_seed", "is_bona_fide_tf"]].copy()
print("source data for Figure 2 (top-25 union):"); display(fig2_src)
"""))
    c.append(code("""
fig, (axA, axB) = plt.subplots(1, 2, figsize=(12.5, 5.8), gridspec_kw={"width_ratios": [1.05, 1]})
cols = [STATUS_COLOR.get(status_map.get(t), NEUTRAL) for t in prim_b.TF]
axA.scatter(prim_b.gremln_dense_rank, prim_b.genie3_dense_rank, c=cols, s=16, alpha=0.75, edgecolor="none")
rho = spearmanr(prim_b.gremln_dense_rank, prim_b.genie3_dense_rank)[0]
axA.set_xlabel("GREmLN dense rank"); axA.set_ylabel("GENIE3 dense rank")
axA.set_title("Rank\\u2013rank, bona-fide common universe (n=%d)" % len(prim_b))
axA.text(0.03, 0.97, "Spearman \\u03c1 = %.2f" % rho, transform=axA.transAxes, va="top", fontsize=10)
for t in PRINCIPAL:
    r = prim_b[prim_b.TF == t]
    if len(r):
        axA.annotate(t, (r.gremln_dense_rank.iloc[0], r.genie3_dense_rank.iloc[0]), fontsize=7.5,
                     xytext=(3, 3), textcoords="offset points")
axA.legend(handles=[Line2D([0], [0], marker="o", color="w", markerfacecolor=SHARED, markersize=8, label="shared top-25"),
                    Line2D([0], [0], marker="o", color="w", markerfacecolor=GRE, markersize=8, label="GREmLN-only top-25"),
                    Line2D([0], [0], marker="o", color="w", markerfacecolor=GEN, markersize=8, label="GENIE3-only top-25"),
                    Line2D([0], [0], marker="o", color="w", markerfacecolor=NEUTRAL, markersize=8, label="other bona-fide TF")],
           loc="lower right", fontsize=7.5)
panel_label(axA, "A")

u = union.sort_values("gremln_dense_rank").reset_index(drop=True); y = np.arange(len(u))
axB.hlines(y, u.gremln_dense_rank, u.genie3_dense_rank, color=NEUTRAL, lw=1, zorder=1)
axB.scatter(u.gremln_dense_rank, y, color=GRE, s=24, zorder=3, label="GREmLN rank")
axB.scatter(u.genie3_dense_rank, y, color=GEN, s=24, zorder=3, label="GENIE3 rank")
labs = [("\\u2605 " if s else "") + ("\\u25B3 " if not b else "") + t
        for t, s, b in zip(u.TF, u.is_BTLA_vs_TCR_seed, u.is_bona_fide_tf)]
axB.set_yticks(y); axB.set_yticklabels(labs, fontsize=6.8); axB.invert_yaxis()
axB.set_xlabel("dense rank (1 = highest priority)")
axB.set_title("Top-25 union, aligned ranks (n=%d)" % len(u))
axB.legend(loc="lower right", fontsize=8)
panel_label(axB, "B")
save_fig(fig, "fig2_primary_rank_comparison", fig2_src,
    "Figure 2. Primary seed-excluded TF prioritisation. (A) Rank\\u2013rank scatter over the bona-fide "
    "common TF universe (n=%d); each point is a TF, coloured by top-25 membership (green shared, blue "
    "GREmLN-only, orange GENIE3-only, grey other); Spearman \\u03c1 annotated; principal candidates "
    "labelled. (B) Aligned dense ranks for the top-25 union: dots joined per TF (blue GREmLN, orange "
    "GENIE3); \\u2605 marks observed BTLA seed TFs, \\u25B3 marks non-bona-fide broader regulators. Ranks, "
    "not raw incomparable scores, are compared." % len(prim_b))
plt.show()
"""))

    c.append(md("## 4. Figure 3 — CRISPRi target-response benchmark"))
    c.append(code("""
cr = D["crispri_all_screen_sensitivity"].copy(); ms = D["crispri_model_summary"].set_index("model")
qcp = D["crispri_primary_qc_passed"].copy()
fig3_src = cr[["TF", "model", "in_screen_8hr", "ontarget_qc_pass", "n_predicted_native",
               "n_confirmed_native_8hr", "frac_native_8hr", "p_native_8hr", "response_direction",
               "validation_status", "frac_top5_8hr", "frac_top10_8hr"]].copy()
print("source data for Figure 3:"); display(fig3_src.head(12))
"""))
    c.append(code("""
fig = plt.figure(figsize=(13.5, 15)); gs = fig.add_gridspec(2, 2, height_ratios=[2.5, 1], hspace=0.16, wspace=0.28)
axA = fig.add_subplot(gs[0, :]); axB = fig.add_subplot(gs[1, 0]); axC = fig.add_subplot(gs[1, 1])

qp = qcp.copy(); qp["label"] = qp["TF"] + " (" + qp["model"].str[:3] + ")"
qp = qp.sort_values(["model", "frac_native_8hr"], ascending=[True, True]).reset_index(drop=True)
yy = np.arange(len(qp))
axA.barh(yy, qp["frac_native_8hr"].astype(float),
         color=[DIR_COLOR.get(d, "#DDDDDD") for d in qp["response_direction"]],
         edgecolor=[GRE if m == "GREmLN" else GEN for m in qp["model"]], lw=1.4)
for i, r in qp.iterrows():
    axA.text(float(r["frac_native_8hr"]) + 0.005, i, "%d/%d" % (int(r["n_confirmed_native_8hr"]), int(r["n_predicted_native"])),
             va="center", fontsize=7.5)
axA.set_yticks(yy); axA.set_yticklabels(qp["label"], fontsize=8); axA.set_ylim(-0.6, len(qp) - 0.4)
axA.set_xlabel("native target-response fraction (confirmed / predicted), Stim8hr")
axA.set_title("A  Candidate-level CRISPRi target response (on-target-QC-passed, in-screen)", loc="left", fontsize=11)
axA.legend(handles=[Line2D([0], [0], marker="s", color="w", markerfacecolor=DIR_COLOR["BTLA_concordant"], markersize=9, label="BTLA-concordant"),
                    Line2D([0], [0], marker="s", color="w", markerfacecolor=DIR_COLOR["anti_concordant"], markersize=9, label="anti-concordant"),
                    Line2D([0], [0], marker="s", color="w", markerfacecolor=DIR_COLOR["mixed"], markersize=9, label="mixed"),
                    Line2D([0], [0], marker="s", color="w", markerfacecolor=DIR_COLOR["no_response"], markersize=9, label="no response"),
                    Line2D([0], [0], marker="s", color="w", markeredgecolor=GRE, markerfacecolor="w", markersize=9, label="GREmLN (edge)"),
                    Line2D([0], [0], marker="s", color="w", markeredgecolor=GEN, markerfacecolor="w", markersize=9, label="GENIE3 (edge)")],
           loc="lower right", fontsize=7)
absent = cr[cr.validation_status.isin(["not_in_screen", "failed_ontarget_qc"])]
txt = "Not scored (shown, not zero):\\n" + "; ".join(sorted(set(absent.TF + " [" + absent.validation_status.str.replace("_", " ") + "]")))
axA.text(1.005, 0.5, txt, transform=axA.transAxes, fontsize=6.6, va="center", ha="left",
         bbox=dict(boxstyle="round", fc="#f7f7f7", ec="#cccccc"))

models = ["GREmLN", "GENIE3"]; means = [ms.loc[m, "mean_frac_native_8hr"] for m in models]
lo = [ms.loc[m, "ci_lo"] for m in models]; hi = [ms.loc[m, "ci_hi"] for m in models]
err = np.array([[means[i] - lo[i] for i in range(2)], [hi[i] - means[i] for i in range(2)]])
axB.bar(models, means, color=[GRE, GEN], width=0.6, yerr=err, capsize=6, edgecolor="#333")
axB.axhline(0, color="#888", lw=0.8)
d_mean = means[0] - means[1]
axB.set_ylabel("mean target-response fraction (QC-passed)")
axB.set_title("B  Aggregate QC-passed comparison", loc="left", fontsize=11)
axB.text(0.5, 0.94, "\\u0394(GREmLN\\u2212GENIE3) = %+.3f\\nsuperiority margin = 0.020" % d_mean,
         transform=axB.transAxes, ha="center", va="top", fontsize=8.5)

budg = pd.DataFrame({"budget": ["native", "top5", "top10"],
    "GREmLN": [ms.loc["GREmLN", "mean_frac_native_8hr"], ms.loc["GREmLN", "mean_frac_top5_8hr"], ms.loc["GREmLN", "mean_frac_top10_8hr"]],
    "GENIE3": [ms.loc["GENIE3", "mean_frac_native_8hr"], ms.loc["GENIE3", "mean_frac_top5_8hr"], ms.loc["GENIE3", "mean_frac_top10_8hr"]]})
xx = np.arange(3); w = 0.38
axC.bar(xx - w / 2, budg["GREmLN"], w, color=GRE, label="GREmLN")
axC.bar(xx + w / 2, budg["GENIE3"], w, color=GEN, label="GENIE3")
axC.set_xticks(xx); axC.set_xticklabels(["native", "top-5", "top-10"])
axC.set_ylabel("mean target-response fraction"); axC.set_title("C  Matched target-budget sensitivity", loc="left", fontsize=11)
axC.legend(fontsize=8)
save_fig(fig, "fig3_crispri_benchmark", fig3_src,
    "Figure 3. Orthogonal CD4 CRISPRi target-response benchmark (Stim8hr). (A) Per-candidate native "
    "target-response fraction for on-target-QC-passed, in-screen TFs; bar fill encodes response "
    "direction, bar edge encodes model (blue GREmLN, orange GENIE3), and k/n denominators are printed; "
    "TFs absent from the screen or failing on-target QC are listed separately, never plotted as zero. "
    "(B) Aggregate QC-passed means with bootstrap 95% CI (5,000 resamples), zero reference and the "
    "0.020 superiority margin. (C) Matched target-budget sensitivity (native / top-5 / top-10 predicted "
    "targets per TF).")
plt.show()
"""))

    c.append(md("## 5. Figure 4 — integrated candidate evidence map"))
    c.append(code("""
ci = D["candidate_integrated_evidence"].copy()
order = {"shared": 0, "GREmLN_specific": 1, "GENIE3_specific": 2}
ci["_o"] = ci["status"].map(order).fillna(3); ci = ci.sort_values(["_o", "gremln_rank"]).reset_index(drop=True)
fig4_src = ci.drop(columns=["_o"]).copy()
print("source data for Figure 4:"); display(fig4_src.head(12))
"""))
    c.append(code("""
def _supp(v): return str(v).lower() not in ("", "nan", "none", "no", "false", "0", "0.0")
TIER = {"strong": 3, "moderate": 2, "weak": 1, "none": 0}
groups = [("model", ["gremln_rank", "genie3_rank"]),
          ("functional", ["frac_native_8hr", "response_direction", "ontarget_qc_pass"]),
          ("independent", ["protein_support", "phosphosite_support", "coip_support"]),
          ("derived", ["transcript_support", "tf_activity_support", "kinase_activity_support", "bionic_support"]),
          ("literature", ["paperclip_evidence_tier"])]
colnames = [x for _, cols in groups for x in cols]
head = {"gremln_rank": "GREmLN\\nrank", "genie3_rank": "GENIE3\\nrank", "frac_native_8hr": "CRISPRi\\nfrac",
        "response_direction": "dir", "ontarget_qc_pass": "QC", "protein_support": "prot",
        "phosphosite_support": "phos", "coip_support": "coIP", "transcript_support": "trans",
        "tf_activity_support": "TFact", "kinase_activity_support": "kin", "bionic_support": "BIONIC",
        "paperclip_evidence_tier": "lit tier"}
nrow, ncol = len(ci), len(colnames)
fig, ax = plt.subplots(figsize=(12.5, max(6, 0.32 * nrow + 2)))
ax.set_xlim(0, ncol); ax.set_ylim(0, nrow); ax.invert_yaxis(); ax.axis("off")
for r in range(nrow):
    row = ci.iloc[r]
    for jc, col in enumerate(colnames):
        v = row[col]; fc = "#FFFFFF"; txt = ""
        if col in ("gremln_rank", "genie3_rank"):
            txt = "" if pd.isna(v) else str(int(v)); fc = "#EFEFEF"
        elif col == "frac_native_8hr":
            fv = float(v) if pd.notna(v) else np.nan
            fc = plt.cm.Blues(min(0.85, 0.15 + 4 * fv)) if np.isfinite(fv) else "#F3F3F3"
            txt = EN if not np.isfinite(fv) else "%.2f" % fv
        elif col == "response_direction":
            fc = DIR_COLOR.get(str(v), "#F3F3F3"); txt = {"BTLA_concordant": "+", "anti_concordant": "\\u2212", "mixed": "~", "no_response": ""}.get(str(v), "")
        elif col == "ontarget_qc_pass":
            fc = "#C8E6C9" if v == True else ("#FFCDD2" if v == False else "#F3F3F3"); txt = "\\u2713" if v == True else ("\\u2717" if v == False else EN)
        elif col == "paperclip_evidence_tier":
            t = TIER.get(str(v), 0); fc = plt.cm.Purples(0.15 + 0.22 * t); txt = str(v) if pd.notna(v) else EN
        else:
            on = _supp(v); fc = "#4D4D4D" if on else "#F0F0F0"; txt = ""
        ax.add_patch(Rectangle((jc, r), 1, 1, fc=fc, ec="white", lw=1.2))
        if txt:
            tc = "white" if (col == "response_direction" and str(v) == "anti_concordant") else "#111"
            ax.text(jc + 0.5, r + 0.5, txt, ha="center", va="center", fontsize=7, color=tc)
    lab = ("\\u2605 " if row.get("is_BTLA_vs_TCR_seed", False) else "") + row["TF"]
    ax.text(-0.2, r + 0.5, lab, ha="right", va="center", fontsize=7.2)
    ax.add_patch(Rectangle((-0.06, r + 0.06), 0.05, 0.88, fc=STATUS_COLOR.get(row["status"], NEUTRAL), ec="none"))
x = 0
for gname, gcols in groups:
    ax.text(x + len(gcols) / 2, -0.7, gname, ha="center", va="center", fontsize=9, fontweight="bold")
    ax.plot([x, x], [0, nrow], color="#999", lw=1.4); x += len(gcols)
ax.plot([x, x], [0, nrow], color="#999", lw=1.4)
for jc, col in enumerate(colnames):
    ax.text(jc + 0.5, -0.15, head[col], ha="center", va="center", fontsize=7)
save_fig(fig, "fig4_integrated_evidence_map", fig4_src,
    "Figure 4. Integrated evidence across the seed-excluded top-25 union (rows; \\u2605 = observed BTLA "
    "seed; left colour bar = shared/green, GREmLN/blue, GENIE3/orange). Columns are grouped into model "
    "prioritisation (ranks), functional validation (CRISPRi native fraction as a blue gradient, "
    "direction +/\\u2212/~, on-target QC \\u2713/\\u2717), independent multi-omics (protein/phospho/co-IP; "
    "filled = present), derived/contextual multi-omics (transcript/TF-activity/kinase/BIONIC), and "
    "literature (Paperclip tier, purple gradient). No cells are combined into a single opaque score.")
plt.show()
"""))

    c.append(md("## 6. Figure 5 — coverage and universe audit"))
    c.append(code("""
fig5_src = fig1_src.copy()
fig, (axF, axH) = plt.subplots(1, 2, figsize=(13, 6), gridspec_kw={"width_ratios": [1.25, 1]})
funnel = [("BTLA panel", n_panel), ("GREmLN-avail seeds", n_gre), ("GENIE3-reach seeds", n_gen),
          ("common seeds", n_cs), ("common regulator universe", n_univ), ("bona-fide TF universe", n_bona),
          ("seed-excluded primary", n_prim), ("top-25 union", n_union), ("in CRISPRi screen", n_screen),
          ("on-target QC passed", n_qc)]
labels = [f[0] for f in funnel]; vals = [f[1] for f in funnel]; yy = np.arange(len(funnel))[::-1]
axF.barh(yy, vals, color="#6BAED6", edgecolor="#2171B5")
for i, v in zip(yy, vals): axF.text(v + max(vals) * 0.01, i, str(v), va="center", fontsize=8)
axF.set_yticks(yy); axF.set_yticklabels(labels, fontsize=8); axF.set_xlabel("count (genes / seeds / TFs)")
axF.set_title("A  Analysis-universe funnel", loc="left", fontsize=11)
crn = D["crispri_all_screen_sensitivity"]
for m, col in (("GREmLN", GRE), ("GENIE3", GEN)):
    vv = crn[crn.model == m]["n_predicted_native"].astype(float).dropna()
    axH.hist(vv, bins=range(0, int(crn["n_predicted_native"].max()) + 2), alpha=0.6, color=col, label=m)
axH.set_xlabel("predicted BTLA targets per TF (native)"); axH.set_ylabel("number of top-25 TFs")
axH.set_title("B  Predicted-target-count distribution", loc="left", fontsize=11); axH.legend(fontsize=8)
save_fig(fig, "fig5_coverage_audit", fig5_src,
    "Figure 5. Coverage and analysis-universe audit. (A) Funnel of exact counts from the 249-gene BTLA "
    "panel through GREmLN-available and GENIE3-reachable seeds, the identical common seed set, the "
    "common regulator universe, the bona-fide TF universe, the seed-excluded primary universe, the "
    "top-25 union, and CRISPRi in-screen / on-target-QC-passed candidates (a funnel, not a proportional "
    "Venn). (B) Distribution of predicted BTLA-target counts per top-25 TF for each model.")
plt.show()
"""))

    c.append(md("## 7. Figure 6 — evidence-tiered shortlist"))
    c.append(code("""
mo = D["multiomics_candidate_evidence"].set_index("TF"); pc = D["paperclip_candidate_evidence"].set_index("TF")
sl = ci.copy()
def _corrob(tf):
    ind = bool(mo.loc[tf, "independent_orthogonal_support"]) if tf in mo.index else False
    lit = str(pc.loc[tf, "paperclip_evidence_tier"]).lower() in ("strong", "moderate") if tf in pc.index else False
    return int(ind) + int(lit)
sl["functional"] = sl["frac_native_8hr"].astype(float).fillna(0.0)
sl["corroboration"] = [ _corrob(t) for t in sl["TF"] ]
fig6_src = sl[["TF", "status", "is_BTLA_vs_TCR_seed", "functional", "corroboration", "crispri_status"]].copy()
print("source data for Figure 6:"); display(fig6_src)
"""))
    c.append(code("""
fig, ax = plt.subplots(figsize=(10.5, 7))
rng = np.random.default_rng(0); texts = []
for _, r in sl.iterrows():
    x = r["functional"]; y = r["corroboration"] + rng.uniform(-0.12, 0.12)
    col = STATUS_COLOR.get(r["status"], NEUTRAL)
    seed = bool(r["is_BTLA_vs_TCR_seed"])
    ax.scatter(x, y, s=70, color=col, edgecolor=(SEED_EDGE if seed else "white"),
               linewidth=(1.8 if seed else 0.8), marker=("D" if seed else "o"), zorder=3)
    if r["TF"] in PRINCIPAL:
        texts.append(ax.text(x, y, r["TF"], fontsize=8))
try:
    from adjustText import adjust_text
    adjust_text(texts, ax=ax, only_move={"text": "xy"}, arrowprops=dict(arrowstyle="-", color="#999", lw=0.5))
except Exception as e:
    print("adjustText skipped:", e)
ax.axvline(0.001, color="#bbb", ls="--", lw=0.8)
ax.set_xlabel("functional target-response evidence (native CRISPRi fraction, Stim8hr)")
ax.set_ylabel("independent + literature corroboration (0\\u20132)")
ax.set_yticks([0, 1, 2]); ax.set_title("Evidence-tiered shortlist of BTLA-response TF candidates")
ax.legend(handles=[Line2D([0], [0], marker="o", color="w", markerfacecolor=SHARED, markersize=10, label="shared"),
                   Line2D([0], [0], marker="o", color="w", markerfacecolor=GRE, markersize=10, label="GREmLN-specific"),
                   Line2D([0], [0], marker="o", color="w", markerfacecolor=GEN, markersize=10, label="GENIE3-specific"),
                   Line2D([0], [0], marker="D", color="w", markerfacecolor="#999", markeredgecolor="k", markersize=10, label="observed seed TF")],
          loc="upper right", fontsize=8)
save_fig(fig, "fig6_candidate_shortlist", fig6_src,
    "Figure 6. Evidence-tiered shortlist. Each candidate is placed by functional target-response "
    "evidence (x, native CRISPRi fraction) and independent+literature corroboration (y, count 0\\u20132: "
    "independent multi-omics and strong/moderate Paperclip). Colour encodes shared (green), "
    "GREmLN-specific (blue) or GENIE3-specific (orange); diamonds outlined in black are observed BTLA "
    "seed TFs (context, not discoveries). Only principal candidates are labelled. No composite score is "
    "used; axes are the raw evidence dimensions.")
plt.show()
"""))


# ------------------------------------------------------------ NB04 tables
def _build_nb4_tables(c):
    c.append(md("## 8. Table 1 — model and benchmark specification"))
    c.append(code("""
t1 = D["model_benchmark_specification"].rename(columns={
    "characteristic": "Characteristic", "implication_for_comparison": "Implication for comparison"})
t1 = t1[["Characteristic", "GENIE3", "GREmLN", "Implication for comparison"]]
print("source data for Table 1:"); display(t1)
out1 = export_table(t1, "table1_model_benchmark_specification",
    "Table 1. Model and benchmark specification",
    ["GENIE3: tree-ensemble co-expression inference on donor-ComBat log1p-CPM CD4 NTC cells.",
     "GREmLN: graph-aware single-cell foundation model, frozen (zero-shot), using the masked GENIE3 graph as prior.",
     "Both models share the identical common seed set and common TF universe; scores are unsigned and compared by rank."])
show_table(out1, "Table 1. Model and benchmark specification")
"""))

    c.append(md("## 9. Table 2 — primary benchmark results"))
    c.append(code("""
ms = D["crispri_model_summary"].set_index("model"); prim_b = primary[primary.is_bona_fide_tf]
rho = spearmanr(prim_b.gremln_dense_rank, prim_b.genie3_dense_rank)[0]
union = D["top25_union_primary"]; shared = int((union.status == "shared").sum())
def _row(metric, gr, ge, diff="", ci="", interp=""):
    return {"Metric": metric, "GREmLN": gr, "GENIE3": ge, "Difference": diff, "95% CI": ci, "Interpretation": interp}
gm, gn = ms.loc["GREmLN"], ms.loc["GENIE3"]
d_mean = gm["mean_frac_native_8hr"] - gn["mean_frac_native_8hr"]
rows = [
    _row("Common BTLA seeds (identical)", n_common_seed, n_common_seed, "0", "", "same seed set (fair)"),
    _row("Bona-fide common TF universe", int(au.is_bona_fide_tf.sum()), int(au.is_bona_fide_tf.sum()), "0", "", "same universe"),
    _row("Spearman rank correlation", "%.3f" % rho, "%.3f" % rho, "", "", "moderate concordance"),
    _row("Top-25 overlap (shared)", shared, shared, "", "", "of 25 each"),
    _row("Present in CRISPRi screen", int(gm["present_in_screen"]), int(gn["present_in_screen"]), "", "", "coverage"),
    _row("Passing on-target QC", int(gm["qc_passed"]), int(gn["qc_passed"]), "", "", "usable perturbations"),
    _row("Included in primary comparison", int(gm["n_in_primary"]), int(gn["n_in_primary"]), "", "", "QC-passed, in-screen"),
    _row("Mean target-response fraction (QC)", "%.3f" % gm["mean_frac_native_8hr"], "%.3f" % gn["mean_frac_native_8hr"],
         "%+.3f" % d_mean, "[%.3f, %.3f] / [%.3f, %.3f]" % (gm["ci_lo"], gm["ci_hi"], gn["ci_lo"], gn["ci_hi"]),
         "within 0.020 margin" if abs(d_mean) < 0.02 else "exceeds margin"),
    _row("Median target-response fraction (QC)", "%.3f" % gm["median_frac_native_8hr"], "%.3f" % gn["median_frac_native_8hr"], "", "", ""),
    _row("BTLA-concordant candidates", int(gm["dir_BTLA_concordant"]), int(gn["dir_BTLA_concordant"]), "", "", "supportive direction"),
    _row("Anti-concordant candidates", int(gm["dir_anti_concordant"]), int(gn["dir_anti_concordant"]), "", "", "opposite direction (not auto-contradictory)"),
    _row("Mixed-direction candidates", int(gm["dir_mixed"]), int(gn["dir_mixed"]), "", "", ""),
    _row("No detected response", int(gm["dir_no_response"]), int(gn["dir_no_response"]), "", "", ""),
    _row("Matched budget: top-5 mean", "%.3f" % gm["mean_frac_top5_8hr"], "%.3f" % gn["mean_frac_top5_8hr"], "", "", "sensitivity"),
    _row("Matched budget: top-10 mean", "%.3f" % gm["mean_frac_top10_8hr"], "%.3f" % gn["mean_frac_top10_8hr"], "", "", "sensitivity"),
    _row("All-screen native mean (sensitivity)", "%.3f" % gm["mean_frac_all_screen_native_8hr"], "%.3f" % gn["mean_frac_all_screen_native_8hr"], "", "", "includes non-QC"),
]
t2 = pd.DataFrame(rows)
print("source data for Table 2:"); display(t2)
out2 = export_table(t2, "table2_primary_benchmark_results", "Table 2. Primary benchmark results",
    ["Target-response fraction = confirmed / predicted BTLA-DEG targets in the Stim8hr CRISPRi arm (adj p < 0.05).",
     "95% CI from 5,000 bootstrap resamples of QC-passed candidates. Superiority margin = 0.020 (pre-specified).",
     "Both scores are unsigned; direction is reported separately and anti-concordance is not automatically contradictory.",
     "En dash (\\u2013) = not applicable."])
show_table(out2, "Table 2. Primary benchmark results")
"""))

    c.append(md("""## 10. Table 3 — top-25 candidate regulators (GREmLN vs GENIE3)

Two ranked lists side by side over the **pre-specified pySCENIC 334-candidate** universe, seed-excluded
primary rankings, unchanged from the canonical package. Shared top-25 candidates are bolded, marked
superscript **S**, and shaded pale green in both columns."""))
    c.append(code("""
# ---- fixed shading + shared-marker ----
SHARE_BG = "#E3F1E1"                    # subtle pale green (shared candidates only)
SUP_S = "\\u02e2"; SIGMA = "\\u03a3"; DASH = "\\u2014"   # superscript s, capital sigma, em dash
def _mbold(sym, suffix=""):             # image-only mathtext: bold symbol + superscript S
    return r"$\\mathbf{%s}^{\\mathrm{S}}$%s" % (sym, suffix)

pr = D["primary_rankings_common_universe"].copy()          # seed-excluded primary (unchanged)
assert not pr["is_BTLA_vs_TCR_seed"].any(), "seeds present in primary ranking"
assert "genie3_n_seed_targets" in pr.columns, "canonical package predates the GENIE3 seed-target count; re-run nb03"
sc = pr.set_index("TF")
u25 = D["top25_union_primary"].set_index("TF")
gr_members = set(u25.index[u25.in_gremln_top25]); g3_members = set(u25.index[u25.in_genie3_top25])
# Membership is taken UNCHANGED from the canonical package. GREmLN ranks by dense seed count with a
# continuous CSLS seed-sum tie-break (this also fixes the top-25 boundary deterministically). GENIE3
# ranks by summed edge weight (near-strict); residual ties are ordered alphabetically.
GR_SORT = (["gremln_dense_rank", "gremln_csls_seed_sum"], [True, False])
assert gr_members == set(pr.sort_values(GR_SORT[0], ascending=GR_SORT[1]).head(25).TF), "GREmLN top-25 differs from canonical"
assert g3_members == set(pr.sort_values("genie3_dense_rank").head(25).TF), "GENIE3 top-25 differs from canonical"
assert len(gr_members) == 25 and len(g3_members) == 25, "top-25 lists are not exactly 25"
gr_ord = pr[pr.TF.isin(gr_members)].sort_values(GR_SORT[0], ascending=GR_SORT[1]).reset_index(drop=True)
g3_ord = pr[pr.TF.isin(g3_members)].sort_values(["genie3_dense_rank", "TF"]).reset_index(drop=True)
shared = gr_members & g3_members
print("Table 3 basis: 25 GREmLN + 25 GENIE3; shared =", len(shared), "->", sorted(shared))

rows, cbg, bold, disp = [], {}, {}, {}
GR_COL, G3_COL = 1, 3
gr_pos = {tf: k + 1 for k, tf in enumerate(gr_ord.TF)}   # displayed ordinal position in the GREmLN half
g3_pos = {tf: k + 1 for k, tf in enumerate(g3_ord.TF)}   # displayed ordinal position in the GENIE3 half
for i in range(25):
    g = gr_ord.TF.iloc[i]; h = g3_ord.TF.iloc[i]
    gs = int(sc.loc[g, "gremln_score"]); gcs = float(sc.loc[g, "gremln_csls_seed_sum"])
    hn = int(sc.loc[h, "genie3_n_seed_targets"]); hw = float(sc.loc[h, "genie3_score"])
    if g in shared:                          # bracket shows this candidate's ORDINAL POSITION in the GENIE3 half
        gr_cell = f"{g}{SUP_S} (G3 {g3_pos[g]})"
        disp[(i, GR_COL)] = _mbold(g, f" (G3 {g3_pos[g]})"); cbg[(i, GR_COL)] = SHARE_BG; bold[(i, GR_COL)] = True
    else:
        gr_cell = g
    if h in shared:                          # bracket shows this candidate's ORDINAL POSITION in the GREmLN half
        g3_cell = f"{h}{SUP_S} (GR {gr_pos[h]})"
        disp[(i, G3_COL)] = _mbold(h, f" (GR {gr_pos[h]})"); cbg[(i, G3_COL)] = SHARE_BG; bold[(i, G3_COL)] = True
    else:
        g3_cell = h
    rows.append({"Position": i + 1, "GREmLN candidate": gr_cell,
                 "GREmLN selection signal": f"{gs} seeds; {SIGMA}CSLS = {gcs:.3f}",
                 "GENIE3 candidate": g3_cell, "GENIE3 selection signal": f"{hn}; {SIGMA}w = {hw:.3f}"})
t3 = pd.DataFrame(rows)
print("source data for Table 3:"); display(t3)
t3_notes = [
    "S, shared between both model top-25 lists; brackets give ordinal position in the other list.",
    "BTLA panel genes were excluded from the primary candidate ranking.",
    "GREmLN uses BTLA seed-neighbour count with summed CSLS as tie-break; GENIE3 uses summed outgoing "
    "edge weight to BTLA seed targets.",
    "Within equal seed-neighbour counts, higher (less negative) summed CSLS ranks higher.",
    "The two model-specific selection signals are not directly comparable."]
export_table(t3, "table3_top25_rankings", "Table 3. Top-25 candidate regulators prioritised by GREmLN and GENIE3",
             t3_notes, cell_bg=cbg, bold_cells=bold, disp=disp, width=1150)

# ---- assertion: every bracketed cross-model position matches the visible position in the other half ----
def _sym(cell): return str(cell).split(SUP_S)[0].split(" (")[0]
gr_seen = {_sym(r["GREmLN candidate"]): r["Position"] for _, r in t3.iterrows()}
g3_seen = {_sym(r["GENIE3 candidate"]): r["Position"] for _, r in t3.iterrows()}
for tf in shared:
    assert g3_pos[tf] == g3_seen[tf] and gr_pos[tf] == gr_seen[tf], f"position map disagrees with table for {tf}"
    gr_cell_tf = t3.loc[t3["GREmLN candidate"].map(_sym) == tf, "GREmLN candidate"].iloc[0]
    g3_cell_tf = t3.loc[t3["GENIE3 candidate"].map(_sym) == tf, "GENIE3 candidate"].iloc[0]
    assert f"(G3 {g3_seen[tf]})" in gr_cell_tf, f"GREmLN-half bracket wrong for {tf}: {gr_cell_tf}"
    assert f"(GR {gr_seen[tf]})" in g3_cell_tf, f"GENIE3-half bracket wrong for {tf}: {g3_cell_tf}"
print("cross-model ordinal-position brackets verified for", len(shared), "shared candidates")
"""))
    c.append(md("### Supplementary Table S1 — full 43-candidate union (ranks and raw scores)"))
    c.append(code("""
sup = u25.reset_index().merge(
    pr[["TF", "gremln_score", "gremln_csls_seed_sum", "genie3_score", "genie3_n_seed_targets"]], on="TF", how="left")
s1 = pd.DataFrame({
    "Candidate": sup.TF, "Model status": sup.status.str.replace("_", " "),
    "GREmLN dense rank": sup.gremln_dense_rank.astype(int),
    "GREmLN seeds (top-100)": sup.gremln_score.astype(int),
    "GREmLN CSLS seed-sum": sup.gremln_csls_seed_sum.round(3),
    "GENIE3 dense rank": sup.genie3_dense_rank.astype(int),
    "GENIE3 linked seed targets": sup.genie3_n_seed_targets.astype(int),
    "GENIE3 summed weight": sup.genie3_score.round(3),
}).sort_values(["GREmLN dense rank", "GREmLN CSLS seed-sum", "GENIE3 dense rank"],
               ascending=[True, False, True]).reset_index(drop=True)
print("Supplementary S1 rows:", len(s1)); display(s1)
export_table(s1, "supplementary_table_s1_full_top25_union",
    "Supplementary Table S1. Full top-25 union of candidate regulators (43 candidates)",
    ["All candidates entering either model's top-25 over the pySCENIC 334-candidate universe.",
     "Model status: shared / GREmLN-specific / GENIE3-specific. Ranks are dense/tie-aware.",
     "GREmLN and GENIE3 raw scores are on different scales and are not directly comparable."], width=1150)
"""))

    c.append(md("""## 11. Table 4 — evidence across each model's top-25

Two matched sub-tables (**4A** GREmLN, **4B** GENIE3) with five columns: `Position`, `Candidate
regulator`, `CRISPRi evidence`, `Literature evidence`, `BTLA experimental evidence`. Every evidence
label is generated from the canonical long-form evidence table (`multiomics_long_evidence.csv`) via an
explicit provenance audit — no labels are hand-typed. For shared candidates the **model-specific**
CRISPRi target set is used in each table. `BTLA experimental evidence` covers the BTLA-engagement
contrasts (**BTLA+TCR vs TCR** and **BTLA+TCR vs UC**); the generic TCR-activation contrast (TCR vs
UC) is excluded and retained only in Supplementary S2."""))
    c.append(md("### Evidence provenance audit (long-form → displayed labels)"))
    c.append(code("""
# ---- unicode marks + colours ----
DAGGER = "\\u2020"; ARR_UP = "\\u2191"; ARR_DN = "\\u2193"
PG = SHARE_BG; GREY = "#EEEEEE"                          # pale green (strongest) / light grey (none/untested)

crx = D["crispri_all_screen_sensitivity"].set_index(["TF", "model"])
pcv = D["paperclip_candidate_evidence"].set_index("TF")
lf = D["multiomics_long_evidence"].copy()

# BTLA experimental evidence covers the BTLA-engagement contrasts (isolate/attribute the BTLA effect);
# the generic TCR-activation contrast (TCR_vs_UC) is excluded and retained only in Supplementary S2.
BTLA_CONTRASTS = ("BTLA_TCR_vs_TCR", "BTLA_TCR_vs_UC"); SUPPORTED = ("moderate", "strong")
LAYER_LABEL = {"proteomics": "Protein" + DAGGER, "phosphoproteomics": "Phosphosite" + DAGGER,
               "coIP": "Co-IP" + DAGGER, "transcriptomics": "Transcript",
               "tf_activity": "Inferred regulon activity", "bionic_gnn": "BIONIC"}
LAYER_ORDER = ["proteomics", "phosphoproteomics", "coIP", "transcriptomics", "tf_activity", "bionic_gnn"]
ORTHO = {"proteomics", "phosphoproteomics", "coIP"}      # orthogonal molecular (dagger; pale green)

def _method(layer):
    if layer == "tf_activity": return "decoupleR ULM (inferred regulon activity)"
    if layer == "bionic_gnn": return "BIONIC GNN network module"
    return "direct measurement"

# provenance audit: one row per candidate-evidence observation, with the inclusion decision recorded
audit_rows = []
for _, r in lf.iterrows():
    layer = str(r["evidence_layer"]); gene = str(r["gene"]); contrast = str(r["contrast"])
    sup = str(r["support_level"]).lower(); direction = str(r["direction"])
    shown = LAYER_LABEL.get(layer); keep = True
    if shown is None:
        keep, reason = False, "layer not shown in main table (" + layer + ")"
    elif sup not in SUPPORTED:
        keep, reason = False, "not significant (support_level=" + sup + ")"
    elif layer == "bionic_gnn":
        reason = "included: BIONIC network module (contrast-free)"
    elif contrast not in BTLA_CONTRASTS:
        keep, reason = False, "contrast not BTLA-specific (" + contrast + "); retained in Supplementary S2"
    else:
        reason = "included: significant at " + contrast
    disp_label = shown if shown else EN
    if keep and layer == "tf_activity":
        disp_label = shown + (" " + ARR_UP if direction == "up" else " " + ARR_DN if direction == "down" else "")
    audit_rows.append({
        "candidate": gene, "model": (u25.loc[gene, "status"] if gene in u25.index else "not_top25"),
        "evidence_layer": layer, "contrast": contrast, "timepoint": r["timepoint"], "condition": r["condition"],
        "direction": direction,
        "effect_or_score": (r["tf_activity_score"] if layer == "tf_activity" else r["effect_size"]),
        "adjusted_p": r["adjusted_p"], "inference_method": _method(layer),
        "regulon_resource": (Path(str(r["source_file"])).name if layer == "tf_activity" else EN),
        "source_file": r["source_file"], "independence": r["independent_or_derived"],
        "displayed_label": (disp_label if keep else EN), "included": keep, "inclusion_reason": reason})
audit = pd.DataFrame(audit_rows)
print("evidence provenance audit rows:", len(audit), "| included:", int(audit["included"].sum()),
      "| candidates with >=1 included:", audit.loc[audit["included"], "candidate"].nunique())
display(audit[audit["included"]].sort_values(["candidate", "evidence_layer", "timepoint"]).reset_index(drop=True))
"""))
    c.append(code("""
# ---- displayed-cell builders (all derived from the audit / canonical CRISPRi) ----
def crispri_cell(tf, model):
    key = (tf, model)
    if key not in crx.index or not bool(crx.loc[key, "in_screen_8hr"]): return "Not in screen"
    r = crx.loc[key]
    if r["ontarget_qc_pass"] != True: return "Failed on-target QC"
    npred = int(r["n_predicted_native"]) if pd.notna(r["n_predicted_native"]) else 0
    nconf = int(r["n_confirmed_native_8hr"]) if pd.notna(r["n_confirmed_native_8hr"]) else 0
    nd = f" ({nconf}/{npred})"; d = str(r["response_direction"])
    return {"BTLA_concordant": f"Detected {DASH} BTLA-concordant" + nd,
            "anti_concordant": f"Detected {DASH} BTLA-anti-concordant" + nd,
            "mixed": f"Detected {DASH} mixed" + nd}.get(d, "No detected response" + nd)

def lit_cell(tf):
    tier = str(pcv.loc[tf, "paperclip_evidence_tier"]).lower() if tf in pcv.index else "none"
    return {"strong": "Strong", "moderate": "Moderate", "weak": "Weak"}.get(tier, "None")

def btla_exp_evidence(tf):
    sub = audit[(audit["candidate"] == tf) & (audit["included"])]
    labels = []
    for layer in LAYER_ORDER:
        ls = sub[sub["evidence_layer"] == layer]
        if not len(ls): continue
        if layer == "tf_activity":
            prim = ls[ls["contrast"] == "BTLA_TCR_vs_TCR"]      # arrow follows the primary contrast if present
            dirs = set((prim if len(prim) else ls)["direction"])
            arrow = (" " + ARR_UP) if dirs == {"up"} else (" " + ARR_DN) if dirs == {"down"} else ""
            labels.append(LAYER_LABEL[layer] + arrow)
        else:
            labels.append(LAYER_LABEL[layer])
    return "; ".join(labels) if labels else "None"

def has_ortho(tf):
    sub = audit[(audit["candidate"] == tf) & (audit["included"])]
    return bool(len(sub[sub["evidence_layer"].isin(ORTHO)]))

CAND_COL, CRIS_COL, LIT_COL, EXP_COL = 1, 2, 3, 4
def build_evidence(model, order_df):
    rows, cbg, bold, disp = [], {}, {}, {}
    for i in range(25):
        tf = order_df.TF.iloc[i]
        if tf in shared:
            cand = f"{tf}{SUP_S}"; disp[(i, CAND_COL)] = _mbold(tf); cbg[(i, CAND_COL)] = PG; bold[(i, CAND_COL)] = True
        else:
            cand = tf
        cris = crispri_cell(tf, model)
        if cris.startswith("Detected"): cbg[(i, CRIS_COL)] = PG                        # any detected response
        elif cris in ("Not in screen", "Failed on-target QC"): cbg[(i, CRIS_COL)] = GREY
        lab = lit_cell(tf)
        if lab == "Strong": cbg[(i, LIT_COL)] = PG
        elif lab == "None": cbg[(i, LIT_COL)] = GREY                                   # moderate/weak -> white
        exp = btla_exp_evidence(tf)
        if has_ortho(tf): cbg[(i, EXP_COL)] = PG                                        # orthogonal molecular
        elif exp == "None": cbg[(i, EXP_COL)] = GREY                                    # transcript/inferred/BIONIC -> white
        rows.append({"Position": i + 1, "Candidate regulator": cand, "CRISPRi evidence": cris,
                     "Literature evidence": lab, "BTLA experimental evidence": exp})
    return pd.DataFrame(rows), cbg, bold, disp

T4_NOTES = [
    "S, shared between the GREmLN and GENIE3 top-25 lists.",
    "Candidate regulators were restricted to the pre-specified pySCENIC human transcription-factor/regulator list.",
    "CRISPRi evidence is model-specific because each model can nominate a different target set for the same candidate.",
    "BTLA-concordant and BTLA-anti-concordant describe response direction and are not automatically supportive or contradictory.",
    "Transcript and inferred regulon activity are corroborating evidence from the BTLA transcriptomic study, not independent validation.",
    "BIONIC is a derived network annotation.",
    DAGGER + ", orthogonal molecular evidence from protein, phosphosite or co-IP data.",
    "Paperclip literature evidence was not used to generate the model rankings.",
    "Inferred regulon activity: decoupleR ULM on the BTLA regulon-activity table (TFactivitiesALL.xlsx), "
    "BTLA-engagement contrasts (BTLA+TCR vs TCR and BTLA+TCR vs UC), 1" + EN + "24 h; " + ARR_UP + "/" + ARR_DN
    + " = increased/decreased inferred activity (following BTLA+TCR vs TCR where available).",
    "Pale green marks the strongest evidence: a detected CRISPRi response (any direction), strong literature, or "
    "orthogonal molecular evidence. Light grey marks not-in-screen, failed QC or no evidence; other cells are white."]

t4a, a_bg, a_bold, a_disp = build_evidence("GREmLN", gr_ord)
t4b, b_bg, b_bold, b_disp = build_evidence("GENIE3", g3_ord)
# identical column widths across 4A/4B so the two evidence profiles are visually comparable
COLW4 = [min(34, max(9, max(len(str(cn)),
             int(t4a[cn].astype(str).map(len).max()), int(t4b[cn].astype(str).map(len).max()))))
         for cn in t4a.columns]
DMAP = {"None": EN}                                        # DOCX-only: replace repeated "None" with an en-dash
print("source data for Table 4A (GREmLN):"); display(t4a)
t4a_out = export_table(t4a, "table4a_gremln_top25_evidence",
    "Table 4A. Evidence across GREmLN's top-25 candidate regulators",
    T4_NOTES, cell_bg=a_bg, bold_cells=a_bold, disp=a_disp, width=1250, col_chars=COLW4, zebra=False, docx_map=DMAP)
"""))
    c.append(code("""
print("source data for Table 4B (GENIE3):"); display(t4b)
t4b_out = export_table(t4b, "table4b_genie3_top25_evidence",
    "Table 4B. Evidence across GENIE3's top-25 candidate regulators",
    T4_NOTES, cell_bg=b_bg, bold_cells=b_bold, disp=b_disp, width=1250, col_chars=COLW4, zebra=False, docx_map=DMAP)
"""))
    c.append(md("### Cross-model evidence summary (descriptive; no weighted score or winner)"))
    c.append(code("""
def summarise(model, order_df):
    d = dict(scr=0, qc=0, use=0, det=0, conc=0, anti=0, lit=0, exp=0, ortho=0)
    for tf in list(order_df.TF):
        key = (tf, model)
        if key in crx.index and bool(crx.loc[key, "in_screen_8hr"]):
            d["scr"] += 1
            if crx.loc[key, "ontarget_qc_pass"] == True:
                d["qc"] += 1; d["use"] += 1
                rd = str(crx.loc[key, "response_direction"])
                if rd in ("BTLA_concordant", "anti_concordant", "mixed"): d["det"] += 1
                if rd == "BTLA_concordant": d["conc"] += 1
                if rd == "anti_concordant": d["anti"] += 1
        if lit_cell(tf) in ("Strong", "Moderate"): d["lit"] += 1
        if btla_exp_evidence(tf) != "None": d["exp"] += 1
        if has_ortho(tf): d["ortho"] += 1
    return d
sg, se = summarise("GREmLN", gr_ord), summarise("GENIE3", g3_ord)
summ = pd.DataFrame([
    {"Evidence summary": "Present in CRISPRi screen", "GREmLN top 25": f"{sg['scr']}/25", "GENIE3 top 25": f"{se['scr']}/25"},
    {"Evidence summary": "Passing on-target QC", "GREmLN top 25": f"{sg['qc']}/25", "GENIE3 top 25": f"{se['qc']}/25"},
    {"Evidence summary": "Target response detected", "GREmLN top 25": f"{sg['det']}/{sg['use']} usable", "GENIE3 top 25": f"{se['det']}/{se['use']} usable"},
    {"Evidence summary": "BTLA-concordant response", "GREmLN top 25": f"{sg['conc']}/{sg['use']} usable", "GENIE3 top 25": f"{se['conc']}/{se['use']} usable"},
    {"Evidence summary": "BTLA-anti-concordant response", "GREmLN top 25": f"{sg['anti']}/{sg['use']} usable", "GENIE3 top 25": f"{se['anti']}/{se['use']} usable"},
    {"Evidence summary": "Strong or moderate literature", "GREmLN top 25": f"{sg['lit']}/25", "GENIE3 top 25": f"{se['lit']}/25"},
    {"Evidence summary": "Any BTLA experimental evidence", "GREmLN top 25": f"{sg['exp']}/25", "GENIE3 top 25": f"{se['exp']}/25"},
    {"Evidence summary": "Orthogonal molecular evidence" + DAGGER, "GREmLN top 25": f"{sg['ortho']}/25", "GENIE3 top 25": f"{se['ortho']}/25"}])
print("Cross-model evidence summary:"); display(summ)
export_table(summ, "table4_evidence_summary", "Table 4. Cross-model evidence summary (GREmLN vs GENIE3 top-25)",
    ["Descriptive comparison of evidence coverage across each model's top-25; no weighted score or winner is implied.",
     "Detected / concordant / anti-concordant responses use QC-passed candidates as the denominator (usable).",
     DAGGER + " Orthogonal molecular evidence = protein, phosphosite or co-IP; BTLA experimental evidence also "
     "includes transcript, inferred regulon activity and BIONIC."],
    width=780)
"""))
    c.append(md("""### Supplementary Table S2 — full long-form candidate evidence

The complete per-observation evidence audit (all layers, contrasts and timepoints for every top-25
union candidate, with the inclusion decision for the main table). Data-only export (CSV/XLSX/HTML)."""))
    c.append(code("""
s2 = audit.sort_values(["candidate", "evidence_layer", "contrast", "timepoint"]).reset_index(drop=True)
print("Supplementary S2 (long-form) rows:", len(s2), "| candidates:", s2.candidate.nunique()); display(s2.head(20))
S2 = "supplementary_table_s2_full_candidate_evidence"
s2.to_csv(TAB / (S2 + ".csv"), index=False)
try: s2.to_excel(TAB / (S2 + ".xlsx"), index=False)
except Exception as e: print("xlsx skipped:", e)
(TAB / (S2 + ".html")).write_text("<h3>Supplementary Table S2. Full long-form candidate evidence</h3>"
    + s2.to_html(index=False, border=0))
(CAP / (S2 + "_caption.txt")).write_text(
    "Supplementary Table S2. Full long-form candidate evidence." + chr(10)
    + "One row per candidate-evidence observation across all layers, contrasts and timepoints; "
      "'included' and 'inclusion_reason' record whether the observation contributes to the main Table 4 "
      "BTLA experimental-evidence column (significant at a BTLA-engagement contrast, BTLA+TCR vs TCR or "
      "BTLA+TCR vs UC, or a BIONIC module)." + chr(10))
print("supplementary S2 written (csv/xlsx/html + caption)")
"""))
    c.append(md("### Table 3/4 quality checks"))
    c.append(code("""
tchecks = {}
tchecks["table3_25_rows"] = len(t3) == 25
tchecks["table4a_25_rows"] = len(t4a) == 25
tchecks["table4b_25_rows"] = len(t4b) == 25
tchecks["rankings_match_canonical"] = (gr_members == set(pr.sort_values(GR_SORT[0], ascending=GR_SORT[1]).head(25).TF)
    and g3_members == set(pr.sort_values("genie3_dense_rank").head(25).TF))
_sh4a = set(t4a.loc[t4a["Candidate regulator"].str.contains(SUP_S), "Candidate regulator"].str.replace(SUP_S, "", regex=False))
_sh4b = set(t4b.loc[t4b["Candidate regulator"].str.contains(SUP_S), "Candidate regulator"].str.replace(SUP_S, "", regex=False))
tchecks["shared_consistent"] = (_sh4a == shared) and (_sh4b == shared)
tchecks["shared_model_specific_crispri"] = all(((tf, "GREmLN") in crx.index and (tf, "GENIE3") in crx.index) for tf in shared)
tchecks["seeds_absent_primary"] = not pr["is_BTLA_vs_TCR_seed"].any()
# failed-QC / not-in-screen must never carry a numeric fraction "(n/N)"
_nofrac = pd.concat([t4a["CRISPRi evidence"], t4b["CRISPRi evidence"]])
tchecks["absent_failed_no_fraction"] = not any(("(" in v) for v in _nofrac
    if v.startswith(("Not in screen", "Failed on-target QC")))
# every displayed BTLA experimental-evidence label traces to >=1 included canonical source row
_inc = set(audit.loc[audit["included"], "candidate"])
tchecks["exp_labels_traceable"] = all(
    (tf in _inc) for t4 in (t4a, t4b)
    for tf, ev in zip(t4["Candidate regulator"].str.replace(SUP_S, "", regex=False), t4["BTLA experimental evidence"])
    if ev != "None")
# every included inferred-regulon-activity row has verified direction/contrast/timepoint/method/regulon source
_ta = audit[(audit["included"]) & (audit["evidence_layer"] == "tf_activity")]
tchecks["regulon_activity_provenance"] = bool(len(_ta)) and (
    _ta["direction"].isin(["up", "down"]).all() and _ta["contrast"].isin(BTLA_CONTRASTS).all()
    and _ta["timepoint"].notna().all() and _ta["inference_method"].notna().all()
    and (_ta["regulon_resource"] != EN).all())
# summary panel equals row-level categories
def _rowcounts(t4):
    det = int(t4["CRISPRi evidence"].str.startswith("Detected").sum())
    scr = int((~t4["CRISPRi evidence"].eq("Not in screen")).sum())
    qc = int((~t4["CRISPRi evidence"].isin(["Not in screen", "Failed on-target QC"])).sum())
    conc = int(t4["CRISPRi evidence"].str.startswith("Detected " + DASH + " BTLA-concordant").sum())
    anti = int(t4["CRISPRi evidence"].str.startswith("Detected " + DASH + " BTLA-anti-concordant").sum())
    lit = int(t4["Literature evidence"].isin(["Strong", "Moderate"]).sum())
    exp = int((t4["BTLA experimental evidence"] != "None").sum())
    return dict(det=det, scr=scr, qc=qc, conc=conc, anti=anti, lit=lit, exp=exp)
_ag = _rowcounts(t4a); _bg2 = _rowcounts(t4b)
tchecks["summary_matches_rows"] = all(_ag[k] == sg[k] for k in _ag) and all(_bg2[k] == se[k] for k in _bg2)
# DOCX/XLSX/CSV/HTML carry identical evidence values (csv==xlsx cell-for-cell; html/docx present and value-complete)
import re as _re
for _n, _tag in [("table4a_gremln_top25_evidence", "4A"), ("table4b_genie3_top25_evidence", "4B")]:
    # keep_default_na=False so the explicit label "None" is not coerced to NaN on readback
    _csv = pd.read_csv(TAB / (_n + ".csv"), keep_default_na=False).astype(str)
    _xls = pd.read_excel(TAB / (_n + ".xlsx"), keep_default_na=False).astype(str)
    _htxt = _re.sub("<[^>]+>", " ", (TAB / (_n + ".html")).read_text())
    _vals_in_html = all((v in _htxt) for v in _csv["BTLA experimental evidence"].unique())
    _docx_ok = (TAB / (_n + ".docx")).exists()
    tchecks["exports_identical_" + _tag] = _csv.equals(_xls) and _vals_in_html and _docx_ok
for k, v in tchecks.items(): print(("PASS" if v else "FAIL"), k)
assert all(tchecks.values()), "Table 3/4 checks failed: " + ", ".join(k for k, v in tchecks.items() if not v)
print(chr(10) + "Table 3/4 quality checks passed.")
"""))


# ------------------------------------------------------------ NB04 index + checks
def _build_nb4_index(c):
    c.append(md("## 12. Asset index and final quality checks"))
    c.append(code("""
figs = [("fig1_study_design", "Figure 1", "Introduction/Methods", "Fair study design; both models share the identical seed set and TF universe."),
        ("fig2_primary_rank_comparison", "Figure 2", "Results", "GREmLN and GENIE3 rankings are moderately correlated; 7/25 shared candidates."),
        ("fig3_crispri_benchmark", "Figure 3", "Results", "Neither model wins CD4 CRISPRi target response beyond the 0.020 margin."),
        ("fig4_integrated_evidence_map", "Figure 4", "Results", "Grouped evidence shows most support is literature/derived, not orthogonal."),
        ("fig5_coverage_audit", "Figure 5", "Results/Methods", "Coverage funnel; GREmLN reach is coverage-limited, not algorithm-limited."),
        ("fig6_candidate_shortlist", "Figure 6", "Results/Discussion", "Few non-seed candidates combine functional and corroborating support.")]
tabs = [("table1_model_benchmark_specification", "Table 1", "Methods", "Specification and the primary fairness limitation."),
        ("table2_primary_benchmark_results", "Table 2", "Results", "Ranking, CRISPRi coverage, primary and matched-budget results."),
        ("table3_top25_rankings", "Table 3", "Results", "Side-by-side top-25 candidate regulators; shared candidates marked S."),
        ("table4a_gremln_top25_evidence", "Table 4A", "Results", "Evidence across GREmLN's top-25 candidate regulators."),
        ("table4b_genie3_top25_evidence", "Table 4B", "Results", "Evidence across GENIE3's top-25 candidate regulators."),
        ("table4_evidence_summary", "Table 4 summary", "Results", "Descriptive evidence-coverage comparison (no winner).")]
supp = [("supplementary_table_s1_full_top25_union", "Supp. Table S1", "Supplement", "Full 43-candidate union with ranks and raw scores."),
        ("supplementary_table_s2_full_candidate_evidence", "Supp. Table S2", "Supplement", "Full 43-candidate evidence across both models.")]

lines = ["# Publication assets index", "",
         "Generated by `notebooks/04_publication_figures_and_tables.ipynb` from the canonical package in",
         "`results/publication_data/`. Figures: pdf/svg/png@600dpi + source_data csv + caption txt.",
         "Tables: csv/xlsx/docx/html + caption txt.", "", "## Figures", ""]
for name, num, sec, msg in figs:
    ok = all((FIG / (name + e)).exists() for e in (".pdf", ".svg", ".png"))
    lines.append("- **%s** (`%s`) - %s - %s - status: %s" % (num, name, sec, msg, "complete" if ok else "MISSING"))
lines += ["", "## Tables", ""]
for name, num, sec, msg in tabs + supp:
    exts = (".csv", ".xlsx", ".html") if name.endswith("s2_full_candidate_evidence") else (".csv", ".xlsx", ".html", ".png", ".pdf", ".svg")
    ok = all((TAB / (name + e)).exists() for e in exts)
    lines.append("- **%s** (`%s`) - %s - %s - status: %s" % (num, name, sec, msg, "complete" if ok else "MISSING"))
lines += ["", "## Outstanding caveats", "",
          "- CRISPRi means differ slightly from the earlier provisional report draft (common-seed + matched-target fix); report to be resynced.",
          "- Held-out seed-TF recovery remains supplementary (excluded from the verdict).",
          "- BTLA-derived assets are gitignored pending data-publication clearance."]
(ASSET / "PUBLICATION_ASSETS_INDEX.md").write_text(chr(10).join(lines) + chr(10))
print("wrote", ASSET / "PUBLICATION_ASSETS_INDEX.md")
display(Markdown(chr(10).join(lines)))
"""))
    c.append(code("""
# visual index: thumbnail of each figure + rendered table
for name, num, sec, msg in figs:
    print("%s  |  %s  |  %s" % (num, name, msg)); display(Image(filename=str(FIG / (name + ".png")), width=680))
for name, num, sec, msg in tabs + supp:
    print("%s  |  %s  |  %s" % (num, name, msg))
    _p = TAB / (name + ".png")
    if _p.exists(): display(Image(filename=str(_p), width=900))
    else: print("(data-only supplementary; see CSV/XLSX/HTML)")
"""))
    c.append(code("""
# ---- final quality checks ----
checks = {}
checks["same_common_seed_set"] = int(seedu.in_common_seed_set.sum()) == n_common_seed and n_common_seed > 0
checks["seeds_absent_from_primary"] = not primary["is_BTLA_vs_TCR_seed"].any()
checks["primary_candidates_classified"] = set(primary.TF).issubset(set(au.TF)) and au.is_bona_fide_tf.notna().all()
checks["absent_failed_not_zero"] = len(crispri[(crispri.validation_status.isin(["not_in_screen", "failed_ontarget_qc"]))
                                               & (crispri.frac_native_8hr.fillna(-1) == 0)]) == 0
checks["model_specific_targets_separate"] = set(crispri.model) == {"GREmLN", "GENIE3"}
# table values match figure/source values
_ms = D["crispri_model_summary"].set_index("model")
_t2mean = float(out2.set_index("Metric").loc["Mean target-response fraction (QC)", "GREmLN"])
checks["table_matches_source"] = abs(_t2mean - round(float(_ms.loc["GREmLN", "mean_frac_native_8hr"]), 3)) < 1e-6
checks["captions_present"] = all((CAP / (n + "_caption.txt")).exists() for n, *_ in figs) and \\
                             all((CAP / (n + "_caption.txt")).exists() for n, *_ in tabs)
checks["assets_from_canonical"] = all((FIG / (n + ".png")).exists() for n, *_ in figs) and \\
                                  all((TAB / (n + ".csv")).exists() for n, *_ in tabs)
for k, v in checks.items():
    print(("PASS" if v else "FAIL"), k)
assert all(checks.values()), "final quality checks failed: " + ", ".join(k for k, v in checks.items() if not v)
print(chr(10) + "All final quality checks passed. Publication assets are in", ASSET)
"""))


# ============================================================ NOTEBOOK 04
def build_nb4():
    c = []
    c.append(md("""
# 04 — Publication figures and tables (BTLA GREmLN vs GENIE3)

Consolidated **publication-asset** notebook. It **does not** re-run GENIE3 or GREmLN and **does not**
redefine the benchmark: every value is loaded from the canonical package frozen by notebook 03 in
`results/publication_data/`. It generates, displays and exports the **six main figures** and **four
main tables**, an asset index, and runs consistency checks.

Assets are written to `results/publication_assets/{figures,tables,captions,source_data}` in multiple
formats (figures: pdf/svg/png@600dpi + source-data csv + caption txt; tables: csv/xlsx/docx/html +
caption txt). Fixed, colour-blind-safe encodings: **GREmLN = blue**, **GENIE3 = orange**,
**shared = green**; observed seed TFs and non-bona-fide (broader-regulator) entries are outlined."""))

    c.append(md("## 0. Setup and input audit"))
    c.append(code("""
import sys, json
from pathlib import Path
import numpy as np, pandas as pd
import matplotlib as mpl, matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Rectangle
from matplotlib.lines import Line2D
from scipy.stats import spearmanr
from IPython.display import display, Image, Markdown

_here = Path.cwd()
_root = next((p for p in [_here, *_here.parents] if (p / "scripts" / "bench_utils.py").exists()), _here.parent)
sys.path.insert(0, str(_root / "scripts"))
import bench_utils as bu

PUB = bu.repo_root() / "results" / "publication_data"
ASSET = bu.repo_root() / "results" / "publication_assets"
FIG, TAB, CAP, SRC = ASSET / "figures", ASSET / "tables", ASSET / "captions", ASSET / "source_data"
for d in (FIG, TAB, CAP, SRC):
    d.mkdir(parents=True, exist_ok=True)

EXPECTED = {
 "model_benchmark_specification": ["characteristic", "GENIE3", "GREmLN", "implication_for_comparison"],
 "common_seed_universe": ["seed_gene", "in_common_seed_set", "gremln_available", "genie3_reachable"],
 "candidate_universe_audit": ["TF", "is_bona_fide_tf", "tf_class", "is_BTLA_vs_TCR_seed"],
 "primary_rankings_common_universe": ["TF", "gremln_dense_rank", "genie3_dense_rank", "is_BTLA_vs_TCR_seed", "is_bona_fide_tf"],
 "seed_inclusive_rankings": ["TF", "gremln_dense_rank", "genie3_dense_rank"],
 "top25_union_primary": ["TF", "status", "in_gremln_top25", "in_genie3_top25", "is_bona_fide_tf"],
 "crispri_primary_qc_passed": ["TF", "model", "frac_native_8hr", "ontarget_qc_pass"],
 "crispri_all_screen_sensitivity": ["TF", "model", "validation_status", "frac_native_8hr", "in_screen_8hr", "response_direction", "n_predicted_native"],
 "crispri_matched_budget_top5": ["TF", "model", "frac_top5_8hr"],
 "crispri_matched_budget_top10": ["TF", "model", "frac_top10_8hr"],
 "crispri_model_summary": ["model", "mean_frac_native_8hr", "ci_lo", "ci_hi", "usable_supportive"],
 "paperclip_candidate_evidence": ["TF", "paperclip_evidence_tier", "paperclip_reviewed"],
 "multiomics_candidate_evidence": ["TF", "independent_orthogonal_support", "derived_contextual_support"],
 "multiomics_long_evidence": ["gene", "evidence_layer", "contrast", "timepoint", "condition", "direction", "support_level"],
 "candidate_integrated_evidence": ["TF", "status", "crispri_status", "frac_native_8hr", "paperclip_evidence_tier"],
}
assert PUB.exists(), f"canonical package missing: {PUB} (run notebook 03 section 12 first)"
D, audit_rows = {}, []
for name, cols in EXPECTED.items():
    p = PUB / (name + ".csv")
    assert p.exists(), f"MISSING canonical input: {p}"
    df = pd.read_csv(p); D[name] = df
    miss = [x for x in cols if x not in df.columns]
    assert not miss, f"{name}.csv missing columns {miss}"
    audit_rows.append({"file": name + ".csv", "rows": len(df), "cols": df.shape[1], "md5": bu.md5(p)})
input_audit = pd.DataFrame(audit_rows)
print("canonical package:", PUB)
print(input_audit.to_string(index=False))
"""))

    c.append(code("""
# ---- stop conditions (fail fast before generating any asset) ----
spec = D["model_benchmark_specification"]; seedu = D["common_seed_universe"]
primary = D["primary_rankings_common_universe"]; au = D["candidate_universe_audit"]
crispri = D["crispri_all_screen_sensitivity"]
n_common_seed = int(seedu["in_common_seed_set"].sum())

assert n_common_seed > 0, "no common seeds"
assert int(seedu["gremln_available"].sum()) >= n_common_seed and int(seedu["genie3_reachable"].sum()) >= n_common_seed
assert str(n_common_seed) in str(spec.loc[spec.characteristic == "common_seed_universe", "GENIE3"].iloc[0]), \\
    "common seed count inconsistent between specification and seed universe (models differ)"
assert not primary["is_BTLA_vs_TCR_seed"].any(), "seeds present in the primary seed-excluded ranking"
_bad = crispri[(crispri.validation_status.isin(["not_in_screen", "failed_ontarget_qc"]))
               & (crispri.frac_native_8hr.fillna(-1) == 0)]
assert len(_bad) == 0, "absent/failed-QC perturbations encoded as zero response"
assert au["is_bona_fide_tf"].notna().all() and au["tf_class"].notna().all(), "unclassified TFs in audit"
assert set(primary["TF"]).issubset(set(au["TF"])), "primary TFs missing bona-fide classification"
# model-specific target sets must be kept separate (both models represented, shared TFs appear twice)
assert set(crispri["model"]) == {"GREmLN", "GENIE3"}, "model-specific CRISPRi rows missing"
print("All input-audit stop-conditions passed. common seeds =", n_common_seed,
      "| bona-fide TFs =", int(au.is_bona_fide_tf.sum()), "of", len(au))
display(input_audit)
"""))

    c.append(md("## 1. Common style and export functions"))
    c.append(code("""
GRE, GEN, SHARED = "#0072B2", "#E69F00", "#009E73"     # colour-blind-safe (Okabe-Ito)
NEUTRAL, SEED_EDGE, NONBONA_EDGE = "#BBBBBB", "#000000", "#D55E00"
STATUS_COLOR = {"shared": SHARED, "GREmLN_specific": GRE, "GENIE3_specific": GEN}
DIR_COLOR = {"BTLA_concordant": "#1B7837", "anti_concordant": "#762A83",
             "mixed": "#B8B8B8", "no_response": "#EEEEEE"}
PRINCIPAL = ["EGR2", "BHLHE40", "JUNB", "IRF8", "AHR", "NFIL3", "PRDM1", "RBPJ", "STAT4", "REL", "TBX21"]
EN = "\\u2013"
mpl.rcParams.update({"figure.dpi": 110, "savefig.dpi": 600, "font.size": 11, "axes.titlesize": 12,
    "axes.labelsize": 11, "axes.spines.top": False, "axes.spines.right": False, "legend.frameon": False,
    "svg.fonttype": "none", "pdf.fonttype": 42, "ps.fonttype": 42, "font.family": "DejaVu Sans"})

def _fmt(v, nd=3):
    if v is None: return EN
    if isinstance(v, float) and not np.isfinite(v): return EN
    if isinstance(v, (int, np.integer)): return str(int(v))
    if isinstance(v, float): return ("%.*f" % (nd, v))
    s = str(v)
    # note: literal "None" is a valid evidence label (Table 4) and is preserved; only true nulls -> EN
    return EN if s.lower() in ("nan", "") else s

def panel_label(ax, s):
    ax.text(-0.10, 1.05, s, transform=ax.transAxes, fontsize=15, fontweight="bold", va="bottom", ha="right")

def save_fig(fig, name, source_df, caption):
    for ext in (".pdf", ".svg", ".png"):
        fig.savefig(FIG / (name + ext), bbox_inches="tight")
    source_df.to_csv(SRC / (name + "_source_data.csv"), index=False)
    (CAP / (name + "_caption.txt")).write_text(caption.strip() + chr(10))
    print("figure saved:", name, "(pdf/svg/png600 + source_data + caption)")

def _html_cell(v, bg=None, bold=False):
    v = str(v)
    st = []
    if bg: st.append("background-color:%s" % bg)
    if bold: st.append("font-weight:bold")
    return '<td style="%s;padding:3px 9px">%s</td>' % (";".join(st), v)

def export_table(df, name, title, footnotes, nd=3, cell_bg=None, bold_cells=None, disp=None, width=980, col_chars=None, zebra=True, docx_map=None):
    # cell_bg: {(row,col): hexcolour}; bold_cells: {(row,col)}; disp: {(row,col): image-only mathtext string}
    # docx_map: {value: replacement} applied ONLY in the DOCX export (e.g. {"None": "-"} for readability);
    #           CSV/XLSX/HTML/PNG retain the explicit values.
    out = df.copy()
    for col in out.columns:
        out[col] = out[col].map(lambda v: _fmt(v, nd))
    cell_bg = cell_bg or {}; bold_cells = bold_cells or set(); disp = disp or {}
    cols = list(out.columns)
    out.to_csv(TAB / (name + ".csv"), index=False)
    try:
        out.to_excel(TAB / (name + ".xlsx"), index=False)
    except Exception as e:
        print("xlsx skipped:", e)
    # HTML with per-cell shading / bold (text labels retained so it reads without colour)
    hrows = ["<tr>" + "".join('<th style="background:#f2f2f2;text-align:left;padding:4px 9px">%s</th>' % c
                              for c in cols) + "</tr>"]
    for i in range(len(out)):
        tds = [_html_cell(out.iloc[i][cols[j]], cell_bg.get((i, j)), (i, j) in bold_cells) for j in range(len(cols))]
        hrows.append("<tr>" + "".join(tds) + "</tr>")
    (TAB / (name + ".html")).write_text("<h3>" + title + "</h3>" + chr(10)
        + '<table style="border-collapse:collapse;font-size:13px">' + chr(10) + chr(10).join(hrows)
        + chr(10) + "</table><hr>" + "<br>".join(footnotes))
    (CAP / (name + "_caption.txt")).write_text(title + chr(10) + chr(10) + chr(10).join(footnotes) + chr(10))
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        def _shade(cell, hexc):
            tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
            sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc.lstrip("#")); tcPr.append(sh)
        doc = Document(); doc.add_heading(title, level=2)
        t = doc.add_table(rows=1, cols=out.shape[1]); t.style = "Table Grid"
        for j, cn in enumerate(cols): t.rows[0].cells[j].text = str(cn)
        for i in range(len(out)):
            cells = t.add_row().cells
            for j, cn in enumerate(cols):
                _val = str(out.iloc[i][cn])
                cells[j].text = (docx_map or {}).get(_val, _val)
                if (i, j) in cell_bg: _shade(cells[j], cell_bg[(i, j)])
                if (i, j) in bold_cells:
                    for p in cells[j].paragraphs:
                        for run in p.runs: run.bold = True
        for fn in footnotes: doc.add_paragraph(fn)
        doc.save(str(TAB / (name + ".docx")))
    except Exception as e:
        print("docx skipped:", e)
    render_table_image(out, name, title, footnotes, cell_bg=cell_bg, bold_cells=bold_cells, disp=disp, col_chars=col_chars, zebra=zebra)
    print("table saved:", name, "(csv/xlsx/docx/html + png/pdf/svg + caption)")
    display(Image(filename=str(TAB / (name + ".png")), width=width))
    return out

def show_table(out, title):
    sty = (out.style.hide(axis="index").set_caption(title).set_table_styles([
        {"selector": "caption", "props": [("font-weight", "bold"), ("font-size", "13px"), ("caption-side", "top"), ("margin-bottom", "6px")]},
        {"selector": "th", "props": [("background-color", "#f2f2f2"), ("text-align", "left"), ("padding", "4px 9px")]},
        {"selector": "td", "props": [("padding", "3px 9px"), ("border-bottom", "1px solid #eee")]}]))
    return sty

def render_table_image(df, name, title, footnotes, cap=30, fs=9, cell_bg=None, bold_cells=None, disp=None, col_chars=None, zebra=True):
    # Publication-styled table rendered to a raster/vector image (booktabs rules, zebra rows,
    # wrapped cells, title above, footnotes below). No browser/HTML dependency.
    # cell_bg {(i,j):colour} shades a body cell (overrides zebra); bold_cells {(i,j)} bolds text;
    # disp {(i,j):str} overrides the drawn text with a verbatim (unwrapped) string, e.g. mathtext.
    # col_chars: fixed per-column wrap widths (chars) so sibling tables share identical geometry.
    import textwrap
    cell_bg = cell_bg or {}; bold_cells = bold_cells or set(); disp = disp or {}
    cols = list(df.columns); n = len(df)
    if col_chars is not None:
        wrapw = [max(9, int(w)) for w in col_chars]
    else:
        raw = [max(len(str(c)), int(df[c].astype(str).map(len).max()) if n else len(str(c))) for c in cols]
        wrapw = [min(cap, max(9, w)) for w in raw]
    tot_w = sum(wrapw)
    frac = [w / tot_w for w in wrapw]; xb = np.concatenate([[0.0], np.cumsum(frac)])
    hdr = [textwrap.fill(str(c), wrapw[j]) for j, c in enumerate(cols)]
    def _txt(i, j):
        if (i, j) in disp: return disp[(i, j)]                    # verbatim (mathtext), no wrap
        return textwrap.fill(str(df.iloc[i][cols[j]]), wrapw[j])
    body = [[_txt(i, j) for j in range(len(cols))] for i in range(n)]
    fnw = [textwrap.fill(fn, max(60, int(tot_w))) for fn in footnotes]
    def _L(row): return max(1, max(s.count(chr(10)) + 1 for s in row))
    hL = _L(hdr); rL = [_L(r) for r in body]; fL = sum(f.count(chr(10)) + 1 for f in fnw)
    top_pad = 1.9; y0 = top_pad + hL; y_end = y0 + sum(rL); total = y_end + 1.1 + fL + 0.4
    figw = min(26, max(6, tot_w * 0.132 + 0.6)); figh = max(2.2, total * 0.30)
    fig, ax = plt.subplots(figsize=(figw, figh)); ax.set_xlim(0, 1); ax.set_ylim(total, 0); ax.axis("off")
    ax.text(0.0, top_pad * 0.45, title, fontsize=fs + 3, fontweight="bold", va="center")
    ax.add_patch(Rectangle((0, top_pad), 1, hL, fc="#eef1f4", ec="none"))
    for j in range(len(cols)):
        ax.text(xb[j] + 0.004, top_pad + hL / 2, hdr[j], fontsize=fs, fontweight="bold", va="center", clip_on=False)
    for yy, lw in ((top_pad, 1.3), (y0, 1.0), (y_end, 1.3)):
        ax.plot([0, 1], [yy, yy], color="#222", lw=lw)
    cy = y0
    for i in range(n):
        h = rL[i]
        if zebra and i % 2 == 1:
            ax.add_patch(Rectangle((0, cy), 1, h, fc="#f7f8fa", ec="none"))
        for j in range(len(cols)):
            if (i, j) in cell_bg:
                ax.add_patch(Rectangle((xb[j], cy), frac[j], h, fc=cell_bg[(i, j)], ec="none"))
        for j in range(len(cols)):
            fw = "bold" if (i, j) in bold_cells else "normal"
            ax.text(xb[j] + 0.004, cy + h / 2, body[i][j], fontsize=fs, va="center", fontweight=fw, clip_on=False)
        cy += h
    fy = y_end + 0.9
    for f in fnw:
        ax.text(0.0, fy, f, fontsize=fs - 1.5, va="top", color="#333", clip_on=False)
        fy += (f.count(chr(10)) + 1) * 0.95
    for ext in (".png", ".pdf", ".svg"):
        fig.savefig(TAB / (name + ext), dpi=600, bbox_inches="tight")
    plt.close(fig)
    print("table image saved:", name + " (.png@600dpi/.pdf/.svg)")

print("style + export helpers ready. GREmLN=blue, GENIE3=orange, shared=green.")
"""))
    _build_nb4_figures(c)
    _build_nb4_tables(c)
    _build_nb4_index(c)
    _write(_nb(c), NB / "04_publication_figures_and_tables.ipynb")


if __name__ == "__main__":
    which = sys.argv[1] if len(sys.argv) > 1 else "all"
    if which in ("01", "all"):
        build_nb1()
    if which in ("02", "all"):
        build_nb2()
    if which in ("03", "all"):
        build_nb3()
    if which in ("04", "all"):
        build_nb4()
